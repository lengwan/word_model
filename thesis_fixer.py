#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
论文格式一键修复引擎
安全原则：只改格式不动内容，生成新文件，逐条记录修复日志
"""

import re
import copy
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

from thesis_checker import (
    FONT_SIZE_MAP, FONT_ALIAS, DEFAULT_RULES, merge_rules,
    SIMSUNG_NAMES, SIMHEI_NAMES, TNR_NAMES,
    W_NS, has_chinese, get_east_asian_font, get_effective_font_size,
    get_effective_line_spacing, get_effective_first_indent,
    get_effective_alignment, pt_to_name,
)

# XML qualified name helper
def qn(tag):
    prefix, tagroot = tag.split(':')
    nsmap = {'w': W_NS}
    return f'{{{nsmap[prefix]}}}{tagroot}'


# ============================================================
# 可修复模块列表（用于 preview 和 UI 展示）
# ============================================================
FIXABLE_MODULES = {'页面设置', '正文格式', '标题层级', '摘要', '目录', '图表规范', '参考文献'}
UNFIXABLE_MODULES = {'页眉页脚', '页码', '编号规范', '单位符号', '内容规范', '其他规范', '封面'}

UNFIXABLE_REASONS = {
    '页眉页脚': '页眉页脚涉及 Word 分节符和域代码，自动修改可能破坏文档结构',
    '页码': '页码格式涉及 Word 域代码，无法通过程序可靠修改',
    '编号规范': '修改图表编号会牵连正文中的引用，需要人工处理',
    '单位符号': '属于内容修改，不是格式修改',
    '内容规范': '属于内容修改（标点、缩写等），不是格式修改',
    '其他规范': '章节结构需要作者自行补充',
    '封面': '封面通常使用内容控件或文本框，自动修改风险较高',
}


class ThesisFixer:
    """论文格式一键修复引擎"""

    def __init__(self, doc_path, issues, rules=None):
        """
        Args:
            doc_path: 原始 .docx 文件路径
            issues: ThesisChecker 输出的 Issue 对象列表
            rules: 格式规则字典（默认使用 DEFAULT_RULES）
        """
        self.doc = Document(doc_path)
        self.issues = issues
        self.rules = rules or copy.deepcopy(DEFAULT_RULES)
        self.fix_log = []     # 已修复记录 [(module, location, description)]
        self.skip_log = []    # 跳过记录 [(module, location, reason)]

    # ── 工具方法 ──

    def _fixed(self, issue, desc=None):
        self.fix_log.append((
            issue.module, issue.location,
            desc or f'{issue.rule} → 已修复'
        ))

    def _skip(self, issue, reason):
        self.skip_log.append((
            issue.module, issue.location,
            reason or issue.rule
        ))

    def _issues_of(self, module_name):
        return [i for i in self.issues if i.module == module_name]

    def _is_cjk(self, text):
        """判断文本是否包含中文"""
        return bool(re.search(r'[\u4e00-\u9fff]', text))

    def _set_east_asian_font(self, run, font_name):
        """设置 run 的东亚字体（直接操作 XML）"""
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:eastAsia'), font_name)

    def _set_font_complete(self, run, cn_font, en_font, size_pt):
        """完整设置 run 的字体：中文字体、英文字体、字号"""
        # 英文/拉丁字体
        run.font.name = en_font
        # 东亚字体
        self._set_east_asian_font(run, cn_font)
        # 字号
        if size_pt:
            run.font.size = Pt(size_pt)

    # ── 预览（不执行修改） ──

    def preview(self):
        """预览修复计划，返回 (可修复数, 跳过数, 详情列表)"""
        fixable = []
        skipped = []
        for issue in self.issues:
            if issue.module in FIXABLE_MODULES and issue.para_index >= 0:
                fixable.append({
                    'module': issue.module,
                    'location': issue.location,
                    'rule': issue.rule,
                    'expected': issue.expected,
                })
            elif issue.module in FIXABLE_MODULES and issue.para_index < 0:
                # 模块可修但无段落定位（如页面设置）
                if issue.module == '页面设置':
                    fixable.append({
                        'module': issue.module,
                        'location': issue.location,
                        'rule': issue.rule,
                        'expected': issue.expected,
                    })
                else:
                    skipped.append({
                        'module': issue.module,
                        'location': issue.location,
                        'reason': '无精确段落定位，跳过',
                    })
            else:
                skipped.append({
                    'module': issue.module,
                    'location': issue.location,
                    'reason': UNFIXABLE_REASONS.get(issue.module, '该类问题暂不支持自动修复'),
                })
        return len(fixable), len(skipped), fixable, skipped

    # ── 执行修复 ──

    def fix_all(self):
        """执行所有可修复项，返回 (fix_log, skip_log)"""
        self._fix_page_setup()
        self._fix_body_text()
        self._fix_headings()
        self._fix_abstract()
        self._fix_toc_title()
        self._fix_captions()
        self._fix_references()

        # 不可修复的模块，统一记录跳过原因
        for issue in self.issues:
            if issue.module in UNFIXABLE_MODULES:
                self._skip(issue, UNFIXABLE_REASONS.get(issue.module, '暂不支持'))

        return self.fix_log, self.skip_log

    # ── 模块1：页面设置 ──

    def _fix_page_setup(self):
        r = self.rules
        issues = self._issues_of('页面设置')
        if not issues:
            return

        target_margins = {
            'top': Cm(r['page']['margin_top_cm']),
            'bottom': Cm(r['page']['margin_bottom_cm']),
            'left': Cm(r['page']['margin_left_cm']),
            'right': Cm(r['page']['margin_right_cm']),
        }

        for section in self.doc.sections:
            changed = False
            if abs(section.top_margin - target_margins['top']) > Cm(0.15):
                section.top_margin = target_margins['top']
                changed = True
            if abs(section.bottom_margin - target_margins['bottom']) > Cm(0.15):
                section.bottom_margin = target_margins['bottom']
                changed = True
            if abs(section.left_margin - target_margins['left']) > Cm(0.15):
                section.left_margin = target_margins['left']
                changed = True
            if abs(section.right_margin - target_margins['right']) > Cm(0.15):
                section.right_margin = target_margins['right']
                changed = True

            # A4 纸张
            a4w, a4h = Cm(21), Cm(29.7)
            if abs(section.page_width - a4w) > Cm(0.5):
                section.page_width = a4w
                changed = True
            if abs(section.page_height - a4h) > Cm(0.5):
                section.page_height = a4h
                changed = True

        if changed:
            for issue in issues:
                self._fixed(issue, '页面设置（页边距/纸张）已修复')

    # ── 模块2：正文格式 ──

    def _fix_body_text(self):
        r = self.rules
        cn_font = r['body']['cn_font']
        en_font = r['body']['en_font']
        target_size = FONT_SIZE_MAP[r['body']['font_size']]
        target_spacing = r['body']['line_spacing']
        target_indent_cm = r['body']['first_indent_char'] * 0.37  # 2字符 ≈ 0.74cm

        paras = self.doc.paragraphs
        fixed_paras = set()

        for issue in self._issues_of('正文格式'):
            idx = issue.para_index
            if idx < 0 or idx >= len(paras):
                self._skip(issue, '段落索引超出范围')
                continue

            para = paras[idx]

            if '字号' in issue.rule or '字体' in issue.rule:
                for run in para.runs:
                    if not run.text.strip():
                        continue
                    if self._is_cjk(run.text):
                        self._set_font_complete(run, cn_font, en_font, target_size)
                    else:
                        run.font.name = en_font
                        run.font.size = Pt(target_size)
                self._fixed(issue, f'字体字号 → {cn_font}/{en_font} {r["body"]["font_size"]}')
                fixed_paras.add(idx)

            elif '行距' in issue.rule:
                para.paragraph_format.line_spacing = target_spacing
                self._fixed(issue, f'行距 → {target_spacing}倍')
                fixed_paras.add(idx)

            elif '缩进' in issue.rule:
                para.paragraph_format.first_line_indent = Cm(target_indent_cm)
                self._fixed(issue, f'首行缩进 → {r["body"]["first_indent_char"]}字符')
                fixed_paras.add(idx)

            else:
                # 未处理的规则类型整体跳过：避免把公式段、代码段、引用块、图题跟随段等
                # 特殊排版强行改成"宋体小四 1.5倍行距 2字符缩进"导致文档结构被破坏
                self._skip(issue, f'未处理的规则类型，请手动调整：{issue.rule}')

    # ── 模块3：标题层级 ──

    def _fix_headings(self):
        r = self.rules
        paras = self.doc.paragraphs

        # 构建各级标题的目标格式
        heading_rules = {}
        for lvl, key in [(1, 'h1'), (2, 'h2'), (3, 'h3')]:
            h = r['headings'][key]
            heading_rules[lvl] = {
                'font': h['font'],
                'size_pt': FONT_SIZE_MAP[h['size']],
                'size_name': h['size'],
                'bold': h.get('bold', True),
            }

        for issue in self._issues_of('标题层级'):
            idx = issue.para_index
            if idx < 0 or idx >= len(paras):
                self._skip(issue, '段落索引超出范围')
                continue

            # 从 issue.location 提取标题级别
            level = None
            m = re.search(r'(\d)级标题', issue.location)
            if m:
                level = int(m.group(1))
            if not level or level not in heading_rules:
                self._skip(issue, f'无法识别标题级别')
                continue

            hr = heading_rules[level]
            para = paras[idx]

            for run in para.runs:
                if not run.text.strip():
                    continue
                self._set_east_asian_font(run, hr['font'])
                run.font.name = hr['font']  # 英文也用同一字体
                run.font.size = Pt(hr['size_pt'])
                run.font.bold = hr['bold']

            self._fixed(issue, f'{level}级标题 → {hr["font"]} {hr["size_name"]} 加粗')

    # ── 模块4：摘要 ──

    def _fix_abstract(self):
        r = self.rules
        paras = self.doc.paragraphs

        cn_title_font = r['abstract']['title_cn_font']
        cn_title_size = FONT_SIZE_MAP[r['abstract']['title_cn_size']]
        en_title_font = r['abstract']['title_en_font']
        en_title_size = FONT_SIZE_MAP[r['abstract']['title_en_size']]

        for issue in self._issues_of('摘要'):
            idx = issue.para_index
            if idx < 0 or idx >= len(paras):
                self._skip(issue, '段落索引超出范围')
                continue

            para = paras[idx]

            if '居中' in issue.rule:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._fixed(issue, '摘要标题 → 居中')

            elif '字号' in issue.rule or '字体' in issue.rule:
                is_english = 'Abstract' in issue.location or 'abstract' in para.text.lower()
                for run in para.runs:
                    if not run.text.strip():
                        continue
                    if is_english:
                        run.font.name = en_title_font
                        run.font.size = Pt(en_title_size)
                    else:
                        self._set_east_asian_font(run, cn_title_font)
                        run.font.name = cn_title_font
                        run.font.size = Pt(cn_title_size)
                self._fixed(issue, '摘要标题字体字号已修复')

            elif '加粗' in issue.rule or '关键词' in issue.rule:
                for run in para.runs:
                    if run.text.strip():
                        run.font.bold = True
                self._fixed(issue, '关键词 → 加粗')

            else:
                self._skip(issue, '摘要问题类型暂不支持自动修复')

    # ── 模块5：目录标题 ──

    def _fix_toc_title(self):
        r = self.rules
        paras = self.doc.paragraphs
        toc_font = r['toc']['title_font']
        toc_size = FONT_SIZE_MAP[r['toc']['title_size']]

        for issue in self._issues_of('目录'):
            idx = issue.para_index
            if idx < 0 or idx >= len(paras):
                self._skip(issue, '段落索引超出范围')
                continue

            para = paras[idx]

            if '居中' in issue.rule:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._fixed(issue, '目录标题 → 居中')

            elif '字号' in issue.rule or '字体' in issue.rule:
                for run in para.runs:
                    if not run.text.strip():
                        continue
                    self._set_east_asian_font(run, toc_font)
                    run.font.name = toc_font
                    run.font.size = Pt(toc_size)
                self._fixed(issue, f'目录标题 → {toc_font} {r["toc"]["title_size"]}')

            else:
                self._skip(issue, '目录问题类型暂不支持自动修复')

    # ── 模块6：图表题注 ──

    def _fix_captions(self):
        r = self.rules
        paras = self.doc.paragraphs
        cap_font = r['caption']['font']
        cap_size = FONT_SIZE_MAP[r['caption']['size']]

        for issue in self._issues_of('图表规范'):
            idx = issue.para_index
            if idx < 0 or idx >= len(paras):
                self._skip(issue, '无段落定位或索引超出范围')
                continue

            if '字号' in issue.rule or '字体' in issue.rule:
                para = paras[idx]
                for run in para.runs:
                    if not run.text.strip():
                        continue
                    if self._is_cjk(run.text):
                        self._set_font_complete(run, cap_font, 'Times New Roman', cap_size)
                    else:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(cap_size)
                self._fixed(issue, f'图表题注 → {cap_font} {r["caption"]["size"]}')
            else:
                self._skip(issue, '图表问题（非字体字号类）暂不支持自动修复')

    # ── 模块7：参考文献 ──

    def _fix_references(self):
        r = self.rules
        paras = self.doc.paragraphs
        # 参考文献用正文字体字号
        cn_font = r['body']['cn_font']
        en_font = r['body']['en_font']
        target_size = FONT_SIZE_MAP[r['body']['font_size']]

        for issue in self._issues_of('参考文献'):
            idx = issue.para_index
            if idx < 0 or idx >= len(paras):
                self._skip(issue, '无段落定位或索引超出范围')
                continue

            if '字号' in issue.rule or '字体' in issue.rule:
                para = paras[idx]
                for run in para.runs:
                    if not run.text.strip():
                        continue
                    if self._is_cjk(run.text):
                        self._set_font_complete(run, cn_font, en_font, target_size)
                    else:
                        run.font.name = en_font
                        run.font.size = Pt(target_size)
                self._fixed(issue, '参考文献字体字号已修复')
            else:
                self._skip(issue, '参考文献问题（非字体字号类）暂不支持自动修复')

    # ── 保存 ──

    def save(self, output_path):
        """保存修复后的文档到新路径"""
        self.doc.save(output_path)
        return output_path

    # ── 修复摘要 ──

    def get_summary(self):
        """返回修复摘要"""
        return {
            'fixed_count': len(self.fix_log),
            'skipped_count': len(self.skip_log),
            'fix_log': self.fix_log,
            'skip_log': self.skip_log,
            'fixed_modules': list(set(m for m, _, _ in self.fix_log)),
            'skipped_modules': list(set(m for m, _, _ in self.skip_log)),
        }
