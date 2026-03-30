#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
硕士毕业论文格式审查工具
用法: py thesis_checker.py <论文.docx>
输出: 格式审查报告.html
"""

import sys
import os
import re
import math
import copy
from dataclasses import dataclass, field
from typing import Optional
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

# ============================================================
# 常量定义
# ============================================================

# 字号映射 (名称 -> pt值)
FONT_SIZE_MAP = {
    '初号': 42, '小初': 36,
    '一号': 26, '小一': 24,
    '二号': 22, '小二': 18,
    '三号': 16, '小三': 15,
    '四号': 14, '小四': 12,
    '五号': 10.5, '小五': 9,
    '六号': 7.5, '小六': 6.5,
    '七号': 5.5, '八号': 5,
}

# pt值反查字号名
PT_TO_NAME = {v: k for k, v in FONT_SIZE_MAP.items()}

# 中文字体名（python-docx 中可能出现的变体）
SIMSUNG_NAMES = {'宋体', 'SimSun', 'simsun', '新宋体', 'NSimSun'}
SIMHEI_NAMES = {'黑体', 'SimHei', 'simhei'}
KAITI_NAMES = {'楷体', 'KaiTi', 'kaiti', '楷体_GB2312'}
TNR_NAMES = {'Times New Roman', 'times new roman', 'TimesNewRomanPSMT'}

# 页边距容差 (EMU)
MARGIN_TOLERANCE = Cm(0.15)  # 允许 1.5mm 误差
A4_WIDTH = Cm(21)
A4_HEIGHT = Cm(29.7)
SIZE_TOLERANCE = Cm(0.5)
TARGET_MARGIN = Cm(2.5)

# Word XML 命名空间
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

# 论文必需章节
REQUIRED_SECTIONS = [
    '符号说明', '摘要', 'abstract', '目录',
    '引言', '材料与方法', '结果与分析', '讨论', '结论',
    '参考文献', '致谢', '攻读学位期间发表论文'
]

# ============================================================
# 默认格式规则（山东农业大学硕士论文）
# ============================================================
DEFAULT_RULES = {
    "school_name": "山东农业大学",
    "degree": "硕士",
    "page": {
        "size": "A4",
        "margin_top_cm": 2.5,
        "margin_bottom_cm": 2.5,
        "margin_left_cm": 2.5,
        "margin_right_cm": 2.5,
    },
    "body": {
        "cn_font": "宋体",
        "en_font": "Times New Roman",
        "font_size": "小四",
        "line_spacing": 1.5,
        "first_indent_char": 2,
    },
    "headings": {
        "h1": {"font": "黑体", "size": "三号", "bold": True},
        "h2": {"font": "黑体", "size": "四号", "bold": True},
        "h3": {"font": "黑体", "size": "小四", "bold": True},
        "numbering": "numeric",
    },
    "abstract": {
        "title_cn_font": "黑体",
        "title_cn_size": "三号",
        "title_en_font": "Times New Roman",
        "title_en_size": "三号",
        "word_count_min": 600,
        "word_count_max": 1500,
        "keyword_count_min": 3,
        "keyword_count_max": 5,
    },
    "toc": {
        "title_font": "黑体",
        "title_size": "三号",
    },
    "caption": {
        "font": "宋体",
        "size": "小五",
        "bilingual": True,
    },
    "header": {
        "odd_page": "山东农业大学硕士学位论文",
        "even_page": "auto",
        "odd_even_different": True,
    },
    "references": {
        "min_count": 100,
        "foreign_ratio": 0.33,
        "cn_before_en": True,
    },
    "cover_fields": ["分类号", "UDC", "学科专业", "指导教师", "论文作者", "培养单位"],
}

# ============================================================
# 字体别名映射 & 工具函数
# ============================================================
FONT_ALIAS = {
    '宋体': SIMSUNG_NAMES,
    '黑体': SIMHEI_NAMES,
    '楷体': KAITI_NAMES,
    'Times New Roman': TNR_NAMES,
}


def _font_match(actual_font, expected_font_name):
    """判断实际字体是否匹配期望字体名"""
    aliases = FONT_ALIAS.get(expected_font_name, {expected_font_name})
    return actual_font in aliases


def merge_rules(custom_rules=None):
    """合并自定义规则和默认规则，custom 覆盖 default，缺失项用默认值"""
    if not custom_rules:
        return copy.deepcopy(DEFAULT_RULES)
    merged = copy.deepcopy(DEFAULT_RULES)
    for key, val in custom_rules.items():
        if isinstance(val, dict) and key in merged and isinstance(merged[key], dict):
            merged[key].update({k: v for k, v in val.items() if v is not None})
        elif val is not None:
            merged[key] = val
    return merged


# ============================================================
# 数据类
# ============================================================

@dataclass
class Issue:
    """单条格式问题"""
    module: str          # 所属检查模块
    severity: str        # error / warning / info
    location: str        # 位置描述
    para_index: int      # 段落序号 (-1 表示非段落)
    text_preview: str    # 文本预览 (截取前40字)
    rule: str            # 违反的规则描述
    expected: str        # 期望值
    actual: str          # 实际值
    source: str          # 'official' 或 'supplement'

    @property
    def source_label(self):
        return {'official': '官方规定', 'supplement': '专业补充',
                'annotation': '批注修订'}.get(self.source, '专业补充')

    @property
    def severity_label(self):
        return {'error': '错误', 'warning': '警告', 'info': '建议'}[self.severity]


# ============================================================
# 辅助函数
# ============================================================

def get_east_asian_font(run):
    """获取 run 的东亚字体名（中文字体）"""
    rPr = run._element.find(f'{{{W_NS}}}rPr')
    if rPr is not None:
        rFonts = rPr.find(f'{{{W_NS}}}rFonts')
        if rFonts is not None:
            ea = rFonts.get(f'{{{W_NS}}}eastAsia')
            if ea:
                return ea
            # 有时中文字体放在 hint="eastAsia" 的 ascii 属性中
            hint = rFonts.get(f'{{{W_NS}}}hint')
            if hint == 'eastAsia':
                return rFonts.get(f'{{{W_NS}}}ascii', None)
    return None


def get_effective_font_name(run):
    """获取 run 的有效字体名（先查直接格式，再查样式）"""
    # 直接格式
    name = run.font.name
    if name:
        return name
    # 样式字体
    style = run._element.getparent()  # w:p
    if style is not None:
        pPr = style.find(f'{{{W_NS}}}pPr')
        if pPr is not None:
            pStyle = pPr.find(f'{{{W_NS}}}pStyle')
            if pStyle is not None:
                return None  # 从样式继承，难以准确获取
    return None


def get_effective_font_size(run, para):
    """获取 run 的有效字号（pt），考虑继承"""
    if run.font.size:
        return run.font.size / 12700  # EMU to pt
    # 尝试从段落样式获取
    if para.style and para.style.font and para.style.font.size:
        return para.style.font.size / 12700
    return None


def get_effective_bold(run, para):
    """获取 run 是否加粗，考虑继承（含XML直接检测）"""
    # 1. run 级别直接设置
    if run.font.bold is not None:
        return run.font.bold
    # 2. 直接检查 XML 中的 <w:b/> 标签（兜底）
    rPr = run._element.find(f'{{{W_NS}}}rPr')
    if rPr is not None:
        b_elem = rPr.find(f'{{{W_NS}}}b')
        if b_elem is not None:
            val = b_elem.get(f'{{{W_NS}}}val', 'true')
            return val.lower() not in ('false', '0', 'off')
    # 3. 段落样式继承
    if para.style and para.style.font and para.style.font.bold is not None:
        return para.style.font.bold
    # 4. 字符样式继承
    if run.style and run.style.font and run.style.font.bold is not None:
        return run.style.font.bold
    return None


def get_effective_alignment(para):
    """获取段落有效对齐方式"""
    if para.alignment is not None:
        return para.alignment
    if para.style and para.style.paragraph_format and para.style.paragraph_format.alignment is not None:
        return para.style.paragraph_format.alignment
    return None


def get_effective_line_spacing(para):
    """获取段落有效行距（倍数）"""
    pf = para.paragraph_format
    if pf.line_spacing is not None:
        if isinstance(pf.line_spacing, (int, float)) and pf.line_spacing < 10:
            return pf.line_spacing
        elif isinstance(pf.line_spacing, int) and pf.line_spacing > 100:
            return pf.line_spacing / 12700 / 12  # EMU -> 行距倍数近似
    # 从样式获取
    if para.style and para.style.paragraph_format:
        spf = para.style.paragraph_format
        if spf.line_spacing is not None:
            if isinstance(spf.line_spacing, (int, float)) and spf.line_spacing < 10:
                return spf.line_spacing
    return None


def get_effective_first_indent(para):
    """获取段落首行缩进（cm）"""
    pf = para.paragraph_format
    if pf.first_line_indent is not None:
        return pf.first_line_indent / 360000  # EMU to cm
    if para.style and para.style.paragraph_format:
        spf = para.style.paragraph_format
        if spf.first_line_indent is not None:
            return spf.first_line_indent / 360000
    return None


def pt_to_name(pt_val):
    """将 pt 值转换为中文字号名"""
    if pt_val is None:
        return '未知'
    # 允许 0.5pt 误差
    for name, pt in FONT_SIZE_MAP.items():
        if abs(pt_val - pt) < 0.6:
            return f'{name}({pt}pt)'
    return f'{pt_val}pt'


def truncate(text, max_len=50):
    """截取文本预览"""
    text = text.replace('\n', ' ').strip()
    if len(text) > max_len:
        return text[:max_len] + '...'
    return text if text else '(空)'


def has_chinese(text):
    """检测文本是否包含中文"""
    return bool(re.search(r'[\u4e00-\u9fff]', text))


def is_heading_paragraph(para):
    """判断段落是否为标题（通过样式名或文本模式）"""
    style_name = para.style.name.lower() if para.style else ''
    if 'heading' in style_name or 'toc' in style_name:
        return True
    text = para.text.strip()
    # 匹配 "1 引言", "2.1 xxx", "3.1.2 xxx" 等
    if re.match(r'^\d+(\.\d+)*\s+\S', text):
        return True
    return False


def get_heading_level(para):
    """获取标题级别: 1=一级, 2=二级, 3=三级及以下, 0=非标题"""
    text = para.text.strip()
    style_name = para.style.name.lower() if para.style else ''

    # 通过样式判断
    m = re.search(r'heading\s*(\d)', style_name)
    if m:
        return min(int(m.group(1)), 3)

    # 通过文本模式判断（允许编号与文字之间有或无空格）
    if re.match(r'^[1-9]\s*[\u4e00-\u9fff]', text) and not re.match(r'^\d+\.\d+', text):
        if len(text) < 30 or re.match(r'^[1-9]\s*(引|前|材料|结果|讨|结论|参考|附|致|概|绪)', text):
            return 1  # "1 引言" or "1引言" or "1前   言"
    if re.match(r'^[1-9]\.\d+[\s\u4e00-\u9fffA-Z]', text) and not re.match(r'^\d+\.\d+\.\d+', text):
        return 2  # "2.1 xxx" or "2.1试验材料" or "2.1 ALV..."
    if re.match(r'^\d+\.\d+\.\d+', text):
        return 3  # "3.1.2 xxx"
    return 0


# ============================================================
# 主检查类
# ============================================================

class ThesisChecker:
    def __init__(self, filepath, thesis_title=None, rules=None):
        self.filepath = filepath
        self.filename = os.path.basename(filepath)
        self.doc = Document(filepath)
        self.issues: list[Issue] = []
        self.scores = {}  # module -> (earned, total)
        self._user_title = thesis_title  # 用户指定的论文题目（优先级最高）
        self.rules = merge_rules(rules)

        # 预处理：识别文档结构
        self._identify_sections()

    def _identify_sections(self):
        """识别文档各部分的段落范围"""
        self.section_ranges = {}
        paras = self.doc.paragraphs
        n = len(paras)

        # 关键词匹配各章节起始位置
        markers = {
            'cover': None,
            'declaration': None,
            'symbols': None,
            'toc': None,
            'abstract_cn': None,
            'abstract_en': None,
            'introduction': None,
            'materials': None,
            'results': None,
            'discussion': None,
            'conclusion': None,
            'references': None,
            'appendix': None,
            'acknowledgement': None,
            'publications': None,
        }

        for i, para in enumerate(paras):
            text = para.text.strip()
            text_lower = text.lower().replace(' ', '')

            if i < 20 and markers['cover'] is None:
                markers['cover'] = 0

            if '原创' in text and '声明' in text:
                markers['declaration'] = i
            elif re.match(r'^符\s*号\s*说\s*明', text):
                markers['symbols'] = i
            elif re.match(r'^目\s*录', text) and len(text) < 10:
                markers['toc'] = i
            elif re.match(r'^中\s*文\s*摘\s*要', text) or text == '摘要' or re.match(r'^摘\s+要', text):
                markers['abstract_cn'] = i
            elif markers['abstract_en'] is None and not has_chinese(text) and (
                text_lower.startswith('abstract') or
                (text_lower == 'abstract') or
                ('abstract' in text_lower and len(text) < 30)
            ):
                markers['abstract_en'] = i
            elif re.match(r'^1\s*(引\s*言|前\s*言|绪\s*论)', text) or re.match(r'^引\s*言', text) or re.match(r'^前\s*言', text):
                markers['introduction'] = i
            elif re.match(r'^2\s*材料与方法', text) or re.match(r'^2\s+材料', text):
                markers['materials'] = i
            elif re.match(r'^3\s*结果与分析', text) or re.match(r'^3\s+结果', text):
                markers['results'] = i
            elif re.match(r'^4\s*讨\s*论', text) or re.match(r'^4\s+讨论', text):
                markers['discussion'] = i
            elif re.match(r'^5\s*结\s*论', text) or re.match(r'^5\s+结论', text):
                markers['conclusion'] = i
            elif re.match(r'^(6\s*)?参\s*考\s*文\s*献', text):
                markers['references'] = i
            elif re.match(r'^(7\s*)?附\s*录', text):
                markers['appendix'] = i
            elif re.match(r'^(8\s*)?致\s*谢', text):
                markers['acknowledgement'] = i
            elif '攻读学位期间' in text or '发表论文' in text or '发表的学术' in text:
                markers['publications'] = i

        # If English abstract not found via 'abstract' keyword, look for the first
        # all-English paragraph after Chinese keywords line (关键词) as the English abstract start.
        if markers['abstract_en'] is None and markers['abstract_cn'] is not None:
            cn_kw_idx = None
            search_end = markers.get('introduction') or min(markers['abstract_cn'] + 80, n)
            for j in range(markers['abstract_cn'] + 1, search_end):
                t = paras[j].text.strip()
                if re.match(r'^关\s*键\s*词', t):
                    cn_kw_idx = j
                    break
            if cn_kw_idx is not None:
                for j in range(cn_kw_idx + 1, search_end):
                    t = paras[j].text.strip()
                    if t and len(t) > 20 and not has_chinese(t) and t[0].isupper():
                        markers['abstract_en'] = j
                        break

        self.markers = markers
        self.total_paras = n

    def _add_issue(self, module, severity, location, para_index, text_preview,
                   rule, expected, actual, source):
        self.issues.append(Issue(
            module=module, severity=severity, location=location,
            para_index=para_index, text_preview=truncate(text_preview),
            rule=rule, expected=expected, actual=actual, source=source
        ))

    # --------------------------------------------------------
    # 检查模块 1: 页面设置
    # --------------------------------------------------------
    def check_page_setup(self):
        """检查页面设置：纸张大小、页边距"""
        module = '页面设置'
        error_count = 0
        total_checks = 0
        r = self.rules

        target_margin_top = Cm(r['page']['margin_top_cm'])
        target_margin_bottom = Cm(r['page']['margin_bottom_cm'])
        target_margin_left = Cm(r['page']['margin_left_cm'])
        target_margin_right = Cm(r['page']['margin_right_cm'])

        for idx, section in enumerate(self.doc.sections):
            sec_label = f'节{idx+1}'

            # 纸张大小
            total_checks += 1
            w, h = section.page_width, section.page_height
            if abs(w - A4_WIDTH) > SIZE_TOLERANCE or abs(h - A4_HEIGHT) > SIZE_TOLERANCE:
                self._add_issue(module, 'error', sec_label, -1, '',
                    '纸张须为A4标准纸', 'A4 (21cm × 29.7cm)',
                    f'{w/360000:.1f}cm × {h/360000:.1f}cm', 'official')
                error_count += 1

            # 页边距
            margins = {
                '上边距': (section.top_margin, target_margin_top, r['page']['margin_top_cm']),
                '下边距': (section.bottom_margin, target_margin_bottom, r['page']['margin_bottom_cm']),
                '左边距': (section.left_margin, target_margin_left, r['page']['margin_left_cm']),
                '右边距': (section.right_margin, target_margin_right, r['page']['margin_right_cm']),
            }
            for name, (val, target, target_cm) in margins.items():
                total_checks += 1
                if val is not None and abs(val - target) > MARGIN_TOLERANCE:
                    self._add_issue(module, 'error', f'{sec_label} {name}', -1, '',
                        f'{name}须为{target_cm}cm', f'{target_cm}cm',
                        f'{val/360000:.2f}cm', 'official')
                    error_count += 1

        total_checks = max(total_checks, 1)
        score = max(0, 10 * (1 - error_count / total_checks))
        self.scores[module] = (round(score, 1), 10)

    # --------------------------------------------------------
    # 检查模块 2: 封面
    # --------------------------------------------------------
    def _extract_sdt_texts(self):
        """从内容控件(sdt)中提取所有文本，返回列表"""
        results = []
        for sdt in self.doc.element.body.iter(f'{{{W_NS}}}sdt'):
            texts = []
            for t in sdt.iter(f'{{{W_NS}}}t'):
                if t.text:
                    texts.append(t.text)
            content = ''.join(texts).strip()
            if content and '点击此处' not in content and '编辑时请删除' not in content:
                results.append(content)
        return results

    def _extract_cover_table_texts(self):
        """从封面表格中提取所有文本"""
        results = {}
        for table in self.doc.tables[:4]:  # 封面表格通常在前4个
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2:
                    key = cells[0].replace('\u3000', '').replace(' ', '').rstrip('：:')
                    val = cells[1].strip() if cells[1].strip() else None
                    if key and val:
                        results[key] = val
        return results

    def check_cover(self):
        """检查封面格式（从段落+内容控件+表格三个来源提取）"""
        module = '封面'
        error_count = 0
        total_checks = 0
        r = self.rules

        # 从三个来源收集封面信息
        sdt_texts = self._extract_sdt_texts()
        table_data = self._extract_cover_table_texts()

        # 所有封面文本合并（用于必填字段检查）
        all_cover_text = ' '.join(sdt_texts) + ' ' + ' '.join(table_data.keys()) + ' ' + ' '.join(str(v) for v in table_data.values())

        # 从段落提取
        paras = self.doc.paragraphs
        cover_end = min(self.markers.get('declaration') or 70, 70)
        for i in range(min(cover_end, len(paras))):
            text = paras[i].text.strip()
            if text:
                all_cover_text += ' ' + text

        # ---- 检查中文题目 ----
        total_checks += 1
        cn_title = None
        # 优先从 sdt 中找（最准确）
        for t in sdt_texts:
            if has_chinese(t) and len(t) > 10 and not any(k in t for k in ['大学','学位','声明','封面','扉页','版权','目录']):
                cn_title = t
                break
        # 从 user_title 回退
        if not cn_title and self._user_title:
            cn_title = self._user_title

        if cn_title:
            # 题目存在，检查字号（sdt 中的字号需要从 XML 获取）
            # 这里只确认题目存在，字号检查从 sdt XML 中难以准确获取，降级为建议
            pass
        else:
            self._add_issue(module, 'warning', '封面', -1, '',
                '未能从段落中检测到中文题目（可能在文本框中）', '应有中文题目',
                '请人工确认题目字体是否为小二号黑体加粗', 'official')
            error_count += 0.5

        # ---- 检查必填字段 ----
        # 从规则中构建字段检查映射
        _cover_field_keywords = {
            '分类号': ['分类号'],
            'UDC': ['UDC'],
            '学科专业': ['专业', '学科', 'Major'],
            '指导教师': ['指导教师', '导师', 'Supervisor'],
            '论文作者': ['论文作者', '研究生', 'Candidate'],
            '培养单位': ['培养单位', '学院', 'College'],
        }
        field_checks = {}
        for field_name in r['cover_fields']:
            if field_name in _cover_field_keywords:
                field_checks[field_name] = _cover_field_keywords[field_name]
            else:
                field_checks[field_name] = [field_name]

        for field_name, keywords in field_checks.items():
            total_checks += 1
            found = False
            # 在 sdt、表格、段落中查找
            for kw in keywords:
                if kw.lower() in all_cover_text.lower():
                    found = True
                    break
            # 在表格 key 中查找
            for table_key in table_data.keys():
                for kw in keywords:
                    if kw in table_key:
                        found = True
                        break

            if not found:
                self._add_issue(module, 'warning', '封面', -1, '',
                    f'封面缺少"{field_name}"信息', f'应包含{field_name}', '未找到', 'official')
                error_count += 0.5

        # ---- 检查封面是否有填写占位符未替换 ----
        total_checks += 1
        placeholders_found = []
        for t in sdt_texts:
            if '点击此处' in t or '填写' in t:
                placeholders_found.append(t[:30])
        if placeholders_found:
            self._add_issue(module, 'warning', '封面', -1, '',
                f'封面有 {len(placeholders_found)} 处占位符未替换',
                '应填写实际内容', f'发现: {placeholders_found[0]}...', 'official')
            error_count += 0.5

        total_checks = max(total_checks, 1)
        score = max(0, 10 * (1 - error_count / total_checks))
        self.scores[module] = (round(score, 1), 10)

    # --------------------------------------------------------
    # 检查模块 3: 摘要
    # --------------------------------------------------------
    def check_abstract(self):
        """检查中英文摘要格式"""
        module = '摘要'
        error_count = 0
        total_checks = 8
        paras = self.doc.paragraphs
        r = self.rules

        # 中文摘要
        cn_idx = self.markers.get('abstract_cn')
        en_idx = self.markers.get('abstract_en')

        exp_cn_title_size = FONT_SIZE_MAP[r['abstract']['title_cn_size']]
        exp_en_title_size = FONT_SIZE_MAP[r['abstract']['title_en_size']]

        if cn_idx is not None:
            # 检查标题格式
            title_para = paras[cn_idx]
            title_text = title_para.text.strip()

            # 标题对齐：应居中
            align = get_effective_alignment(title_para)
            if align != WD_ALIGN_PARAGRAPH.CENTER:
                self._add_issue(module, 'error', f'第{cn_idx+1}段(中文摘要标题)', cn_idx,
                    title_text, '中文摘要标题须居中', '居中对齐',
                    '非居中', 'official')
                error_count += 1

            # 标题字号
            for run in title_para.runs:
                if run.text.strip():
                    size_pt = get_effective_font_size(run, title_para)
                    if size_pt and abs(size_pt - exp_cn_title_size) > 1:
                        self._add_issue(module, 'error', f'第{cn_idx+1}段(中文摘要标题)', cn_idx,
                            title_text, f'中文摘要标题须为{r["abstract"]["title_cn_size"]}',
                            f'{r["abstract"]["title_cn_size"]}({exp_cn_title_size}pt)',
                            pt_to_name(size_pt), 'official')
                        error_count += 1

                    ea_font = get_east_asian_font(run)
                    if ea_font and not _font_match(ea_font, r['abstract']['title_cn_font']):
                        self._add_issue(module, 'error', f'第{cn_idx+1}段(中文摘要标题)', cn_idx,
                            title_text, f'中文摘要标题须为{r["abstract"]["title_cn_font"]}',
                            r['abstract']['title_cn_font'],
                            ea_font, 'official')
                        error_count += 1
                    break

            # 检查关键词
            end_idx = en_idx if en_idx else min(cn_idx + 80, len(paras))
            kw_found = False
            for j in range(cn_idx + 1, end_idx):
                text = paras[j].text.strip()
                if re.match(r'^关\s*键\s*词', text):
                    kw_found = True
                    # 检查"关键词"是否加粗
                    if paras[j].runs:
                        first_run = paras[j].runs[0]
                        bold = get_effective_bold(first_run, paras[j])
                        if bold is False:
                            self._add_issue(module, 'error', f'第{j+1}段(关键词)', j,
                                text, '"关键词"三字须加粗', '加粗',
                                '未加粗', 'supplement')
                            error_count += 1

                    # 检查关键词分隔符
                    kw_content = re.sub(r'^关\s*键\s*词\s*[：:]?\s*', '', text)
                    if kw_content:
                        if '；' not in kw_content and ';' not in kw_content:
                            self._add_issue(module, 'warning', f'第{j+1}段(关键词)', j,
                                text, '关键词之间应用分号分隔', '使用"；"分隔',
                                '未检测到分号', 'supplement')
                            error_count += 0.5

                        # 检查关键词数量
                        kws = re.split(r'[；;]', kw_content)
                        kws = [k.strip() for k in kws if k.strip()]
                        kw_min = r['abstract']['keyword_count_min']
                        kw_max = r['abstract']['keyword_count_max']
                        if len(kws) < kw_min or len(kws) > kw_max:
                            self._add_issue(module, 'warning', f'第{j+1}段(关键词)', j,
                                text, f'关键词数量应为{kw_min}-{kw_max}个',
                                f'{kw_min}-{kw_max}个',
                                f'{len(kws)}个', 'official')
                            error_count += 0.5
                    break

            if not kw_found:
                self._add_issue(module, 'error', '中文摘要区域', cn_idx, '',
                    '未找到关键词', '应有"关键词："行', '未找到', 'official')
                error_count += 1
        else:
            self._add_issue(module, 'error', '全文', -1, '',
                '未找到中文摘要', '应包含中文摘要', '未找到', 'official')
            error_count += 3

        # 英文摘要
        if en_idx is not None:
            title_para = paras[en_idx]
            title_text = title_para.text.strip()

            # 检查对齐
            align = get_effective_alignment(title_para)
            if align != WD_ALIGN_PARAGRAPH.CENTER:
                self._add_issue(module, 'error', f'第{en_idx+1}段(Abstract标题)', en_idx,
                    title_text, 'Abstract标题须居中', '居中对齐',
                    '非居中', 'official')
                error_count += 1

            # 检查字号
            for run in title_para.runs:
                if run.text.strip():
                    size_pt = get_effective_font_size(run, title_para)
                    if size_pt and abs(size_pt - exp_en_title_size) > 1:
                        self._add_issue(module, 'error', f'第{en_idx+1}段(Abstract标题)', en_idx,
                            title_text, f'Abstract标题须为{r["abstract"]["title_en_size"]}',
                            f'{r["abstract"]["title_en_size"]}({exp_en_title_size}pt)',
                            pt_to_name(size_pt), 'official')
                        error_count += 1
                    font_name = run.font.name
                    if font_name and not _font_match(font_name, r['abstract']['title_en_font']):
                        self._add_issue(module, 'error', f'第{en_idx+1}段(Abstract标题)', en_idx,
                            title_text, f'Abstract标题须为{r["abstract"]["title_en_font"]}',
                            r['abstract']['title_en_font'], font_name, 'official')
                        error_count += 1
                    break

            # 检查 Keywords
            search_end = min(en_idx + 80, self.markers.get('introduction') or len(paras))
            kw_en_found = False
            for j in range(en_idx + 1, search_end):
                text = paras[j].text.strip()
                if text.lower().startswith('keyword'):
                    kw_en_found = True
                    if paras[j].runs:
                        bold = get_effective_bold(paras[j].runs[0], paras[j])
                        if bold is False:
                            self._add_issue(module, 'warning', f'第{j+1}段(Keywords)', j,
                                text, '"Keywords"须加粗', '加粗',
                                '未加粗', 'supplement')
                            error_count += 0.5
                    break
            if not kw_en_found:
                self._add_issue(module, 'warning', '英文摘要区域', en_idx, '',
                    '未找到Keywords', '应有Keywords行', '未找到', 'official')
                error_count += 0.5
        else:
            self._add_issue(module, 'error', '全文', -1, '',
                '未找到英文摘要', '应包含Abstract', '未找到', 'official')
            error_count += 2

        score = max(0, 12 * (1 - error_count / total_checks))
        self.scores[module] = (round(score, 1), 12)

    # --------------------------------------------------------
    # 检查模块 4: 目录
    # --------------------------------------------------------
    def check_toc(self):
        """检查目录格式"""
        module = '目录'
        error_count = 0
        total_checks = 3
        paras = self.doc.paragraphs
        r = self.rules

        exp_toc_title_size = FONT_SIZE_MAP[r['toc']['title_size']]

        toc_idx = self.markers.get('toc')
        if toc_idx is None:
            # 检查 sdt 内容控件和 TOC 域代码中是否有目录
            sdt_texts = self._extract_sdt_texts()
            has_toc_in_sdt = any('目' in t and '录' in t for t in sdt_texts)
            # 检查 Word TOC 域代码
            has_toc_field = bool(self.doc.element.body.findall(
                f'.//{{{W_NS}}}fldSimple[@{{{W_NS}}}instr]'))
            if has_toc_in_sdt or has_toc_field:
                self.scores[module] = (6, 8)
                return
            self._add_issue(module, 'error', '全文', -1, '',
                '未找到目录', '应包含目录', '未找到', 'official')
            self.scores[module] = (0, 8)
            return

        title_para = paras[toc_idx]

        # 标题格式：居中
        align = get_effective_alignment(title_para)
        if align != WD_ALIGN_PARAGRAPH.CENTER:
            self._add_issue(module, 'error', f'第{toc_idx+1}段(目录标题)', toc_idx,
                title_para.text, '目录标题须居中', '居中',
                '非居中', 'supplement')
            error_count += 1

        for run in title_para.runs:
            if run.text.strip():
                size_pt = get_effective_font_size(run, title_para)
                if size_pt and abs(size_pt - exp_toc_title_size) > 1:
                    self._add_issue(module, 'error', f'第{toc_idx+1}段(目录标题)', toc_idx,
                        title_para.text, f'目录标题须为{r["toc"]["title_size"]}',
                        f'{r["toc"]["title_size"]}({exp_toc_title_size}pt)',
                        pt_to_name(size_pt), 'supplement')
                    error_count += 1

                ea_font = get_east_asian_font(run)
                if ea_font and not _font_match(ea_font, r['toc']['title_font']):
                    self._add_issue(module, 'error', f'第{toc_idx+1}段(目录标题)', toc_idx,
                        title_para.text, f'目录标题须为{r["toc"]["title_font"]}',
                        r['toc']['title_font'],
                        ea_font, 'supplement')
                    error_count += 1
                break

        score = max(0, 8 * (1 - error_count / total_checks))
        self.scores[module] = (round(score, 1), 8)

    # --------------------------------------------------------
    # 检查模块 5: 正文格式
    # --------------------------------------------------------
    def check_body_text(self):
        """检查正文段落字体、字号、行距、首行缩进"""
        module = '正文格式'
        paras = self.doc.paragraphs
        r = self.rules

        exp_body_size = FONT_SIZE_MAP[r['body']['font_size']]
        exp_line_spacing = r['body']['line_spacing']
        exp_indent_char = r['body']['first_indent_char']
        # 缩进范围：基于字符数计算（每字符约0.37-0.42cm），留一定容差
        indent_min = exp_indent_char * 0.25
        indent_max = exp_indent_char * 0.6

        # 确定正文范围 — only start from introduction (not abstract)
        start = self.markers.get('introduction')
        if start is None:
            start = self.markers.get('materials') or self.markers.get('results') or 0
        end = self.markers.get('references') or len(paras)

        error_count = 0
        checked_count = 0
        font_errors = 0
        size_errors = 0
        spacing_errors = 0
        indent_errors = 0
        max_report_per_type = 999  # 逐条报告，不限数量

        for i in range(start, min(end, len(paras))):
            para = paras[i]
            text = para.text.strip()
            if not text or len(text) < 5:
                continue

            # 跳过标题段落
            level = get_heading_level(para)
            if level > 0:
                continue

            # 跳过图表标注行（图题、表题、注释等有独立格式规范）
            if re.match(r'^(图|表|Fig|Table|Figure)\s*\.?\s*\d', text):
                continue
            if re.match(r'^(注：|注意|Note|Source|[A-Z]{1,2}$)', text):
                continue
            # 跳过英文图表标题行（可能不以数字开头）
            if re.match(r'^(Fig\.|Figure|Table)\s', text, re.IGNORECASE):
                continue

            checked_count += 1
            section = self._get_section_label(i)
            loc = f'第{i+1}段 [{section}]'

            # 检查行距
            line_sp = get_effective_line_spacing(para)
            if line_sp is not None and abs(line_sp - exp_line_spacing) > 0.1:
                spacing_errors += 1
                if spacing_errors <= max_report_per_type:
                    self._add_issue(module, 'error', loc, i, text,
                        f'正文行距须为{exp_line_spacing}倍', f'{exp_line_spacing}倍行距',
                        f'{line_sp:.2f}倍', 'official')

            # 检查首行缩进
            indent_cm = get_effective_first_indent(para)
            if indent_cm is not None:
                if indent_cm < indent_min or indent_cm > indent_max:
                    indent_errors += 1
                    if indent_errors <= max_report_per_type:
                        self._add_issue(module, 'warning', loc, i, text,
                            f'正文须首行缩进{exp_indent_char}字符',
                            f'缩进{exp_indent_char}字符(约0.74-0.85cm)',
                            f'{indent_cm:.2f}cm', 'supplement')
            elif indent_cm is None and checked_count <= 3:
                # 前几段如果没有缩进信息，给个提醒
                pass

            # 检查字体字号（抽样检查 runs）
            for run in para.runs:
                run_text = run.text.strip()
                if not run_text:
                    continue

                # 字号检查
                size_pt = get_effective_font_size(run, para)
                if size_pt and abs(size_pt - exp_body_size) > 0.6:
                    size_errors += 1
                    if size_errors <= max_report_per_type:
                        self._add_issue(module, 'error', loc, i, text,
                            f'正文字号须为{r["body"]["font_size"]}号',
                            f'{r["body"]["font_size"]}({exp_body_size}pt)',
                            pt_to_name(size_pt), 'official')
                    break

                # 中文字体检查
                if has_chinese(run_text):
                    ea_font = get_east_asian_font(run)
                    if ea_font and not _font_match(ea_font, r['body']['cn_font']):
                        font_errors += 1
                        if font_errors <= max_report_per_type:
                            self._add_issue(module, 'warning', loc, i, text,
                                f'正文中文须用{r["body"]["cn_font"]}', r['body']['cn_font'],
                                ea_font, 'official')
                        break
                else:
                    font_name = run.font.name
                    en_font_aliases = FONT_ALIAS.get(r['body']['en_font'], {r['body']['en_font']})
                    cn_font_aliases = FONT_ALIAS.get(r['body']['cn_font'], {r['body']['cn_font']})
                    if font_name and font_name not in en_font_aliases and font_name not in cn_font_aliases:
                        font_errors += 1
                        if font_errors <= max_report_per_type:
                            self._add_issue(module, 'warning', loc, i, text,
                                f'正文英文须用{r["body"]["en_font"]}',
                                r['body']['en_font'], font_name, 'official')
                        break

        total_errors = font_errors + size_errors + spacing_errors + indent_errors

        # 如果超出报告上限，添加汇总信息
        for label, cnt in [('字体', font_errors), ('字号', size_errors),
                           ('行距', spacing_errors), ('缩进', indent_errors)]:
            if cnt > max_report_per_type:
                self._add_issue(module, 'info', '汇总', -1, '',
                    f'共发现 {cnt} 处{label}问题', f'此处仅展示前{max_report_per_type}条',
                    f'总计{cnt}处', 'official' if label != '缩进' else 'supplement')

        checked_count = max(checked_count, 1)
        error_rate = min(total_errors / checked_count, 1.0)
        score = max(0, 20 * (1 - error_rate))
        self.scores[module] = (round(score, 1), 20)

    # --------------------------------------------------------
    # 检查模块 6: 标题层级
    # --------------------------------------------------------
    def check_headings(self):
        """检查标题字体、字号"""
        module = '标题层级'
        paras = self.doc.paragraphs
        start = self.markers.get('introduction') or 0
        end = self.markers.get('references') or len(paras)
        r = self.rules

        error_count = 0
        total_headings = 0

        # 期望值（从规则读取）
        expected = {}
        for lvl, key in [(1, 'h1'), (2, 'h2'), (3, 'h3')]:
            h = r['headings'][key]
            exp_size = FONT_SIZE_MAP[h['size']]
            exp_fonts = FONT_ALIAS.get(h['font'], {h['font']})
            expected[lvl] = (h['size'], exp_size, exp_fonts, h['font'])

        for i in range(start, min(end, len(paras))):
            para = paras[i]
            level = get_heading_level(para)
            if level == 0:
                continue

            total_headings += 1
            text = para.text.strip()
            exp_name, exp_pt, exp_fonts, exp_font_name = expected[level]

            for run in para.runs:
                if not run.text.strip():
                    continue

                # 字号检查
                size_pt = get_effective_font_size(run, para)
                if size_pt and abs(size_pt - exp_pt) > 1:
                    self._add_issue(module, 'error', f'第{i+1}段({level}级标题)', i,
                        text, f'{level}级标题须为{exp_name}',
                        f'{exp_name}({exp_pt}pt)', pt_to_name(size_pt), 'official')
                    error_count += 1

                # 字体检查
                ea_font = get_east_asian_font(run)
                if ea_font and ea_font not in exp_fonts:
                    self._add_issue(module, 'error', f'第{i+1}段({level}级标题)', i,
                        text, f'{level}级标题须为{exp_font_name}', exp_font_name,
                        ea_font, 'official')
                    error_count += 1

                break  # 只检查第一个有文字的 run

        total_headings = max(total_headings, 1)
        score = max(0, 12 * (1 - error_count / (total_headings * 2)))
        self.scores[module] = (round(score, 1), 12)

    # --------------------------------------------------------
    # 检查模块 7: 图表规范
    # --------------------------------------------------------
    def _get_section_label(self, para_idx):
        """根据段落索引返回所在章节标签，如 '第2章 材料与方法'"""
        # 向前搜索最近的一级标题
        paras = self.doc.paragraphs
        for j in range(para_idx, -1, -1):
            level = get_heading_level(paras[j])
            if level == 1:
                title = paras[j].text.strip()[:20]
                return title
        return '封面/前置页'

    def _check_caption_format(self, module, para, para_idx, text, caption_type, total_checks):
        """检查图题/表题的字号、字体，逐条报告"""
        r = self.rules
        EXPECTED_SIZE = FONT_SIZE_MAP[r['caption']['size']]
        EXPECTED_NAME = f'{r["caption"]["size"]}({EXPECTED_SIZE}pt)'

        section = self._get_section_label(para_idx)
        loc = f'第{para_idx+1}段 [{section}]({caption_type})'

        for run in para.runs:
            if not run.text.strip():
                continue
            # 字号
            size_pt = get_effective_font_size(run, para)
            if size_pt and abs(size_pt - EXPECTED_SIZE) > 1.0:
                self._add_issue(module, 'warning', loc, para_idx,
                    text, f'{caption_type}字号应为{EXPECTED_NAME}',
                    EXPECTED_NAME, pt_to_name(size_pt), 'supplement')
            # 中文字体
            if has_chinese(run.text):
                ea_font = get_east_asian_font(run)
                if ea_font and not _font_match(ea_font, r['caption']['font']):
                    self._add_issue(module, 'warning', loc, para_idx,
                        text, f'{caption_type}中文应为{r["caption"]["font"]}',
                        r['caption']['font'], ea_font, 'supplement')
            break

    def check_figures_tables(self):
        """检查图表格式"""
        module = '图表规范'
        paras = self.doc.paragraphs
        error_count = 0
        total_checks = 0

        start = self.markers.get('introduction') or 0
        end = self.markers.get('references') or len(paras)

        fig_pattern = re.compile(r'^(图|Fig\.?|Figure)\s*(\d[\d\-\.]*)')
        tab_pattern = re.compile(r'^(表|Table)\s*(\d[\d\-\.]*)')

        fig_numbers = []
        tab_numbers = []

        for i in range(start, min(end, len(paras))):
            text = paras[i].text.strip()

            # 图题检查
            fig_m = fig_pattern.match(text)
            if fig_m:
                total_checks += 1
                fig_numbers.append((fig_m.group(2), i))

                # 图题应有中英文对照 — check adjacent paragraphs for paired caption
                fig_is_paired = False
                if has_chinese(text) and not re.search(r'[A-Za-z]{3,}', text):
                    if i + 1 < min(end, len(paras)):
                        next_text = paras[i + 1].text.strip()
                        if re.match(r'^(Fig\.?|Figure)\s*\d', next_text, re.IGNORECASE):
                            fig_is_paired = True
                    if not fig_is_paired:
                        self._add_issue(module, 'warning', f'第{i+1}段(图题)', i,
                            text, '图题须中英文对照', '包含英文说明',
                            '仅有中文', 'official')
                        error_count += 1
                elif not has_chinese(text) and re.search(r'[A-Za-z]', text):
                    if i - 1 >= start:
                        prev_text = paras[i - 1].text.strip()
                        if re.match(r'^图\s*\d', prev_text):
                            fig_is_paired = True
                    if not fig_is_paired:
                        self._add_issue(module, 'warning', f'第{i+1}段(图题)', i,
                            text, '图题须中英文对照', '包含中文说明',
                            '仅有英文', 'official')
                        error_count += 1

                # 图题字体字号检查（应为五号=10.5pt）
                self._check_caption_format(module, paras[i], i, text, '图题', total_checks)
                total_checks += 1

            # 表题检查
            tab_m = tab_pattern.match(text)
            if tab_m:
                total_checks += 1
                tab_numbers.append((tab_m.group(2), i))

                # 表题字体字号检查
                self._check_caption_format(module, paras[i], i, text, '表题', total_checks)
                total_checks += 1

                # 表题应有中英文对照 — check adjacent paragraphs for paired caption
                tab_is_paired = False
                if has_chinese(text) and not re.search(r'[A-Za-z]{3,}', text):
                    if i + 1 < min(end, len(paras)):
                        next_text = paras[i + 1].text.strip()
                        if re.match(r'^Table\s*\d', next_text, re.IGNORECASE):
                            tab_is_paired = True
                    if not tab_is_paired:
                        self._add_issue(module, 'warning', f'第{i+1}段(表题)', i,
                            text, '表题须中英文对照', '包含英文说明',
                            '仅有中文', 'official')
                        error_count += 1
                elif not has_chinese(text) and re.search(r'[A-Za-z]', text):
                    # English only — check if previous paragraph is Chinese pair
                    if i - 1 >= start:
                        prev_text = paras[i - 1].text.strip()
                        if re.match(r'^表\s*\d', prev_text):
                            tab_is_paired = True
                    if not tab_is_paired:
                        pass  # English table caption following Chinese is acceptable

        # 检查表格是否使用三线表（通过XML检查边框），逐个表格报告
        for tbl_idx, table in enumerate(self.doc.tables):
            total_checks += 1
            tbl = table._tbl
            tblBorders = tbl.find(f'{{{W_NS}}}tblPr/{{{W_NS}}}tblBorders')

            # 获取表格标识信息
            header_cells = []
            if table.rows and table.rows[0].cells:
                header_cells = [c.text.strip()[:15] for c in table.rows[0].cells[:3]]
            header_preview = ' | '.join(c for c in header_cells if c)[:50]
            # 找表格前面最近的表题
            tbl_name = f'第{tbl_idx+1}个表格'
            # 在段落中搜索该表格前的表号
            tbl_elem = table._tbl
            prev = tbl_elem.getprevious()
            while prev is not None:
                if prev.tag == f'{{{W_NS}}}p':
                    p_text = ''.join(t.text or '' for t in prev.iter(f'{{{W_NS}}}t')).strip()
                    if re.match(r'^(表|Table)\s*\d', p_text):
                        tbl_name = p_text[:30]
                        break
                    if p_text and len(p_text) > 5:
                        break  # 不是表题，停止搜索
                prev = prev.getprevious()

            if tblBorders is not None:
                left = tblBorders.find(f'{{{W_NS}}}left')
                right = tblBorders.find(f'{{{W_NS}}}right')
                has_side_borders = False
                for side in [left, right]:
                    if side is not None:
                        val = side.get(f'{{{W_NS}}}val', 'none')
                        if val not in ('none', 'nil', ''):
                            has_side_borders = True

                if has_side_borders:
                    self._add_issue(module, 'warning',
                        f'{tbl_name}', -1,
                        header_preview,
                        f'表格应采用三线表格式（无左右边框）— {tbl_name}',
                        '三线表（上下粗线+中间细线，无左右线）',
                        '检测到左右边框', 'supplement')
                    error_count += 1

        total_checks = max(total_checks, 1)
        score = max(0, 10 * (1 - error_count / total_checks))
        self.scores[module] = (round(score, 1), 10)

    # --------------------------------------------------------
    # 检查模块 8: 页眉页脚
    # --------------------------------------------------------
    def _get_thesis_title(self):
        """获取论文中文题目（用户指定 > sdt提取 > 封面段落 > 偶数页眉）"""
        if self._user_title:
            return self._user_title
        # 优先从 sdt 内容控件提取
        for sdt in self.doc.element.body.iter(f'{{{W_NS}}}sdt'):
            texts = []
            for t in sdt.iter(f'{{{W_NS}}}t'):
                if t.text:
                    texts.append(t.text)
            content = ''.join(texts).strip()
            if content and has_chinese(content) and len(content) > 10:
                if not any(k in content for k in ['点击此处','大学','声明','封面','扉页','版权','目录','摘要']):
                    return content
        # 从封面段落提取
        paras = self.doc.paragraphs
        cover_end = min(self.markers.get('declaration') or 30,
                        self.markers.get('abstract_cn') or 30, 50)
        skip_kw = ['大学', '学位', '论文提交', '声明', '分类号', '密级',
                    '学号', '授予', '研究生', 'UDC', '答辩', '学科', '导师',
                    '年月', '委员', '专业学位', 'Shandong', 'Agricultural',
                    'Thesis', 'Degree', 'June', 'College', '代码', '指导',
                    '研究方向', '学院', '签名', '日期', '保留']
        for i in range(min(cover_end, len(paras))):
            text = paras[i].text.strip()
            # 跳过含制表符的表单行（如"密级：\t学号："）
            if '\t' in paras[i].text:
                continue
            if has_chinese(text) and len(text) > 10 and not any(k in text for k in skip_kw):
                return text

        # 封面可能用文本框存题目，尝试从偶数页页眉获取
        for section in self.doc.sections:
            even_header = section.even_page_header
            if not even_header.is_linked_to_previous:
                even_text = ''.join(p.text.strip() for p in even_header.paragraphs)
                if even_text and len(even_text) > 5:
                    return even_text
        return None

    def _is_even_odd_headers_enabled(self):
        """检查文档是否开启了'奇偶页不同页眉'设置"""
        settings = self.doc.settings.element
        eaoh = settings.find(f'{{{W_NS}}}evenAndOddHeaders')
        if eaoh is None:
            return False
        # 如果元素存在但无 val 属性，默认为 true
        val = eaoh.get(f'{{{W_NS}}}val', 'true')
        return val.lower() not in ('false', '0', 'off')

    def check_headers_footers(self):
        """检查页眉页脚（奇数页：学校名；偶数页：论文题目）"""
        module = '页眉页脚'
        error_count = 0
        total_checks = 0
        r = self.rules

        ODD_HEADER_EXPECTED = r['header']['odd_page']
        thesis_title = self._get_thesis_title()
        even_odd_enabled = self._is_even_odd_headers_enabled()

        # 如果未开启奇偶页不同，先报告这个问题
        if not even_odd_enabled and r['header']['odd_even_different']:
            total_checks += 1
            self._add_issue(module, 'error', '文档设置', -1, '',
                '未开启"奇偶页不同页眉"（规范要求奇数页为学校名、偶数页为论文题目）',
                '开启奇偶页不同页眉',
                '当前所有页面共用同一页眉', 'official')
            error_count += 1

        for idx, section in enumerate(self.doc.sections):
            # 跳过第一节（封面/声明，通常不设页眉）
            if idx == 0:
                continue

            # ---- 默认页眉（未开启奇偶页时=所有页面，开启后=奇数页）----
            total_checks += 1
            header = section.header
            if not header.is_linked_to_previous:
                default_text = ''.join(p.text.strip() for p in header.paragraphs)
                header_label = '奇数页页眉' if even_odd_enabled else '页眉(所有页面)'

                if not default_text:
                    self._add_issue(module, 'error', f'节{idx+1} {header_label}', -1, '',
                        f'{header_label}不应为空', ODD_HEADER_EXPECTED,
                        '页眉为空', 'official')
                    error_count += 1
                elif default_text != ODD_HEADER_EXPECTED:
                    self._add_issue(module, 'warning', f'节{idx+1} {header_label}', -1,
                        default_text,
                        f'{"奇数页" if even_odd_enabled else ""}页眉内容不符合规范，请核对学校要求',
                        ODD_HEADER_EXPECTED, default_text, 'official')
                    error_count += 0.5

                # 页眉居中检查
                for p in header.paragraphs:
                    if p.text.strip():
                        align = get_effective_alignment(p)
                        if align is not None and align != WD_ALIGN_PARAGRAPH.CENTER:
                            self._add_issue(module, 'warning', f'节{idx+1} {header_label}', -1,
                                p.text.strip(), '页眉应居中', '居中',
                                '非居中', 'supplement')
                            error_count += 0.5
                        break

            # ---- 偶数页页眉（仅在开启奇偶页不同时检查）----
            if even_odd_enabled:
                total_checks += 1
                even_header = section.even_page_header
                if not even_header.is_linked_to_previous:
                    even_text = ''.join(p.text.strip() for p in even_header.paragraphs)
                    if not even_text:
                        self._add_issue(module, 'error', f'节{idx+1} 偶数页页眉', -1, '',
                            '偶数页页眉不应为空（应为论文题目）',
                            thesis_title or '论文题目', '页眉为空', 'official')
                        error_count += 1
                    elif thesis_title and even_text != thesis_title:
                        self._add_issue(module, 'warning', f'节{idx+1} 偶数页页眉', -1,
                            even_text, '偶数页页眉应为论文题目',
                            truncate(thesis_title, 40), truncate(even_text, 40), 'official')
                        error_count += 0.5

        if total_checks == 0:
            total_checks = 1
            self._add_issue(module, 'error', '全文', -1, '',
                '未检测到有效页眉设置', '应设置奇偶页页眉',
                '无页眉', 'official')
            error_count += 1

        total_checks = max(total_checks, 1)
        score = max(0, 8 * (1 - error_count / total_checks))
        self.scores[module] = (round(score, 1), 8)

    # --------------------------------------------------------
    # 检查模块 9: 参考文献
    # --------------------------------------------------------
    def check_references(self):
        """检查参考文献格式和数量"""
        module = '参考文献'
        error_count = 0
        total_checks = 4
        paras = self.doc.paragraphs
        r = self.rules

        ref_start = self.markers.get('references')
        if ref_start is None:
            self._add_issue(module, 'error', '全文', -1, '',
                '未找到参考文献章节', '应包含参考文献', '未找到', 'official')
            self.scores[module] = (0, 5)
            return

        # 确定参考文献结束位置
        ref_end = len(paras)
        for key in ['appendix', 'acknowledgement', 'publications']:
            idx = self.markers.get(key)
            if idx and idx > ref_start:
                ref_end = min(ref_end, idx)
                break

        # 收集参考文献条目
        ref_entries = []
        numbered_refs = 0
        cn_refs = []
        en_refs = []

        for i in range(ref_start + 1, ref_end):
            text = paras[i].text.strip()
            if not text or len(text) < 10:
                continue

            ref_entries.append((i, text))

            # 检查是否有编号
            if re.match(r'^\[\d+\]', text) or re.match(r'^\d+[\.\)]\s', text):
                numbered_refs += 1

            # 分类中英文
            if has_chinese(text):
                cn_refs.append(i)
            else:
                en_refs.append(i)

        # 检查数量
        total_refs = len(ref_entries)
        ref_min = r['references']['min_count']
        if total_refs < ref_min:
            self._add_issue(module, 'error', '参考文献', ref_start, '',
                f'硕士论文参考文献应不少于{ref_min}篇', f'≥{ref_min}篇',
                f'{total_refs}篇', 'official')
            error_count += 1

        # 检查外文比例
        foreign_ratio = r['references']['foreign_ratio']
        if total_refs > 0:
            en_ratio = len(en_refs) / total_refs
            if en_ratio < foreign_ratio:
                self._add_issue(module, 'warning', '参考文献', ref_start, '',
                    f'外文文献应占{foreign_ratio:.0%}以上',
                    f'≥{int(total_refs * foreign_ratio)}篇外文',
                    f'{len(en_refs)}篇外文({en_ratio:.0%})', 'official')
                error_count += 0.5

        # 检查是否有编号（不应编号）
        if numbered_refs > total_refs * 0.3:
            self._add_issue(module, 'error', '参考文献', ref_start, '',
                '参考文献不应编号，首行顶格写', '无编号',
                f'发现{numbered_refs}条带编号', 'official')
            error_count += 1

        # 检查排序：中文在前英文在后
        if r['references']['cn_before_en'] and cn_refs and en_refs:
            last_cn = max(cn_refs)
            first_en = min(en_refs)
            if first_en < last_cn:
                # 检查是否存在中英文交错
                self._add_issue(module, 'warning', '参考文献', ref_start, '',
                    '参考文献应中文在前，英文在后', '中文→英文',
                    '中英文交错排列', 'official')
                error_count += 0.5

        score = max(0, 5 * (1 - error_count / total_checks))
        self.scores[module] = (round(score, 1), 5)

    # --------------------------------------------------------
    # 检查模块 10: 结构完整性
    # --------------------------------------------------------
    def check_structure(self):
        """检查论文结构是否完整"""
        module = '其他规范'
        error_count = 0
        total_checks = 6

        # 检查必需章节
        section_checks = {
            'abstract_cn': ('中文摘要', 'official'),
            'abstract_en': ('英文摘要', 'official'),
            'introduction': ('引言/前言', 'official'),
            'materials': ('材料与方法', 'official'),
            'results': ('结果与分析', 'official'),
            'discussion': ('讨论', 'official'),
            'conclusion': ('结论', 'official'),
            'references': ('参考文献', 'official'),
            'acknowledgement': ('致谢', 'official'),
            'publications': ('攻读学位期间发表论文', 'official'),
        }

        # Collect sdt texts to search for markers not found in paragraphs
        sdt_texts = self._extract_sdt_texts()
        sdt_combined = ' '.join(sdt_texts)

        # Keywords to search in sdt for each marker
        sdt_keywords = {
            'abstract_cn': ['摘要', '中文摘要'],
            'abstract_en': ['Abstract', 'ABSTRACT'],
            'introduction': ['引言', '前言', '绪论'],
            'materials': ['材料与方法', '材料'],
            'results': ['结果与分析', '结果'],
            'discussion': ['讨论'],
            'conclusion': ['结论'],
            'references': ['参考文献'],
            'acknowledgement': ['致谢'],
            'publications': ['攻读学位期间', '发表论文', '发表的学术'],
        }

        missing = []
        for key, (name, source) in section_checks.items():
            if self.markers.get(key) is None:
                # Before reporting missing, check if it exists in sdt content
                found_in_sdt = False
                for kw in sdt_keywords.get(key, []):
                    if kw in sdt_combined:
                        found_in_sdt = True
                        break
                if found_in_sdt:
                    continue  # Found in sdt, not truly missing
                missing.append(name)
                self._add_issue(module, 'error', '论文结构', -1, '',
                    f'缺少"{name}"章节', f'应包含{name}', '未找到', source)
                error_count += 1

        # 检查是否分章写（硕士论文不允许）
        paras = self.doc.paragraphs
        chapter_pattern = re.compile(r'^第[一二三四五六七八九十\d]+章')
        for i, para in enumerate(paras):
            if chapter_pattern.match(para.text.strip()):
                self._add_issue(module, 'error', f'第{i+1}段', i,
                    para.text.strip(), '硕士论文不能按章书写', '使用数字编号(1, 2, 3...)',
                    '检测到"第X章"格式', 'official')
                error_count += 1
                break

        total_checks = max(len(section_checks) + 1, 1)
        score = max(0, 5 * (1 - error_count / total_checks))
        self.scores[module] = (round(score, 1), 5)

    # --------------------------------------------------------
    # 检查模块 11: 图表编号规范
    # --------------------------------------------------------
    def check_numbering(self):
        """检查图号、表号的连续性、一致性、中英文对应"""
        module = '编号规范'
        error_count = 0
        total_checks = 0
        paras = self.doc.paragraphs

        start = self.markers.get('introduction') or 0
        end = len(paras)  # 含附录

        # 收集所有图号和表号
        cn_figs = []   # (编号字符串, 段落索引, 原文)
        en_figs = []
        cn_tabs = []
        en_tabs = []

        fig_cn_pat = re.compile(r'^图\s*(\d[\d\-\.]*)')
        fig_en_pat = re.compile(r'^Fig\.?\s*(\d[\d\-\.]*)', re.IGNORECASE)
        tab_cn_pat = re.compile(r'^表\s*(\d[\d\-\.]*)')
        tab_en_pat = re.compile(r'^Table\s*(\d[\d\-\.]*)', re.IGNORECASE)
        # 检测缺失编号（如"表--"、"Table -"）
        missing_num_pat = re.compile(r'^(图|表|Fig|Table)\s*[-—]{1,}')

        for i in range(start, end):
            text = paras[i].text.strip()
            if not text:
                continue

            m = fig_cn_pat.match(text)
            if m:
                cn_figs.append((m.group(1), i, text))
                continue
            m = fig_en_pat.match(text)
            if m:
                en_figs.append((m.group(1), i, text))
                continue
            m = tab_cn_pat.match(text)
            if m:
                cn_tabs.append((m.group(1), i, text))
                continue
            m = tab_en_pat.match(text)
            if m:
                en_tabs.append((m.group(1), i, text))
                continue

            # 检测缺失编号
            if missing_num_pat.match(text):
                total_checks += 1
                self._add_issue(module, 'error', f'第{i+1}段', i, text,
                    '图/表编号缺失（出现"--"占位符）', '应有完整编号',
                    '编号缺失', 'annotation')
                error_count += 1

        # --- 图号格式一致性 ---
        total_checks += 1
        if cn_figs:
            has_space = sum(1 for num, idx, txt in cn_figs if re.match(r'^图\s+\d', txt))
            no_space = sum(1 for num, idx, txt in cn_figs if re.match(r'^图\d', txt))
            if has_space > 0 and no_space > 0:
                self._add_issue(module, 'warning', '全文图号', -1, '',
                    '图号格式不一致：部分用"图 X"(有空格)，部分用"图X"(无空格)',
                    '全文统一为"图X"或"图 X"',
                    f'"图 X"格式{has_space}处, "图X"格式{no_space}处', 'annotation')
                error_count += 1

        # --- 图号连续性 ---
        total_checks += 1
        if cn_figs:
            nums = []
            for num_str, idx, txt in cn_figs:
                # 纯数字编号（非章节制）
                try:
                    nums.append((int(num_str), idx, txt))
                except ValueError:
                    pass  # 章节制编号如 "3-1"，单独处理

            if nums:
                nums.sort(key=lambda x: x[0])
                for j in range(len(nums) - 1):
                    curr_n, curr_i, curr_t = nums[j]
                    next_n, next_i, next_t = nums[j + 1]
                    if next_n - curr_n > 1:
                        self._add_issue(module, 'error',
                            f'图{curr_n}→图{next_n}', curr_i, curr_t,
                            f'图号不连续：图{curr_n}后应为图{curr_n+1}，实际为图{next_n}',
                            f'图{curr_n+1}', f'图{next_n}（跳号）', 'supplement')
                        error_count += 1
                    elif next_n == curr_n:
                        self._add_issue(module, 'error',
                            f'第{next_i+1}段', next_i, next_t,
                            f'图号重复：图{curr_n}出现多次',
                            '唯一编号', f'图{curr_n}重复', 'supplement')
                        error_count += 1

        # --- 中英文图号对应 ---
        total_checks += 1
        cn_fig_nums = [num for num, idx, txt in cn_figs]
        en_fig_nums = [num for num, idx, txt in en_figs]
        for num in cn_fig_nums:
            if num not in en_fig_nums:
                # 找到对应的中文图所在行
                para_idx = next((idx for n, idx, t in cn_figs if n == num), -1)
                self._add_issue(module, 'warning', f'图{num}', para_idx, '',
                    f'中文图{num}缺少对应的英文Fig.{num}',
                    f'Fig.{num}', '未找到', 'official')
                error_count += 0.5
        for num in en_fig_nums:
            if num not in cn_fig_nums:
                para_idx = next((idx for n, idx, t in en_figs if n == num), -1)
                self._add_issue(module, 'warning', f'Fig.{num}', para_idx, '',
                    f'英文Fig.{num}缺少对应的中文图{num}',
                    f'图{num}', '未找到', 'official')
                error_count += 0.5

        # --- Fig 编号与图编号数字是否匹配 ---
        total_checks += 1
        # 配对相邻的中英文图题
        for cn_num, cn_idx, cn_txt in cn_figs:
            # 找紧跟其后的 Fig
            if cn_idx + 1 < len(paras):
                next_text = paras[cn_idx + 1].text.strip()
                m = fig_en_pat.match(next_text)
                if m and m.group(1) != cn_num:
                    self._add_issue(module, 'error',
                        f'第{cn_idx+1}-{cn_idx+2}段', cn_idx, cn_txt,
                        f'中英文图号不匹配：图{cn_num}对应的Fig编号为{m.group(1)}',
                        f'Fig.{cn_num}', f'Fig.{m.group(1)}', 'supplement')
                    error_count += 1

        # --- 表号连续性（章节编号制） ---
        total_checks += 1
        if cn_tabs:
            # 按章节分组检查
            chapter_tabs = {}
            for num_str, idx, txt in cn_tabs:
                parts = re.split(r'[-\.]', num_str)
                if len(parts) == 2:
                    ch, seq = parts
                    chapter_tabs.setdefault(ch, []).append((int(seq), idx, txt))

            for ch, items in chapter_tabs.items():
                items.sort(key=lambda x: x[0])
                for j in range(len(items) - 1):
                    curr_seq, curr_i, curr_t = items[j]
                    next_seq, next_i, next_t = items[j + 1]
                    if next_seq - curr_seq > 1:
                        self._add_issue(module, 'warning',
                            f'表{ch}-{curr_seq}→表{ch}-{next_seq}', curr_i, curr_t,
                            f'表号不连续：表{ch}-{curr_seq}后应为表{ch}-{curr_seq+1}',
                            f'表{ch}-{curr_seq+1}', f'表{ch}-{next_seq}（跳号）', 'supplement')
                        error_count += 0.5

        total_checks = max(total_checks, 1)
        score = max(0, 8 * (1 - error_count / max(total_checks, 1)))
        self.scores[module] = (round(score, 1), 8)

    # --------------------------------------------------------
    # 检查模块 12: 单位与符号规范
    # --------------------------------------------------------
    def check_units_symbols(self):
        """检查单位规范、化学式、数值单位间距——逐条报告"""
        module = '单位符号'
        paras = self.doc.paragraphs
        start = self.markers.get('introduction') or 0
        end = self.markers.get('references') or len(paras)

        error_count = 0
        total_checks = 0

        unit_rules = [
            (re.compile(r'(?<![a-zA-Z])rpm(?![a-zA-Z])'),
             '转速单位应使用r/min，不用rpm', 'r/min', 'annotation'),
            (re.compile(r'(?<![a-zA-Z/])(?<!\d)ml(?![a-zA-Z])'),
             '毫升应写mL（L大写）', 'mL', 'annotation'),
            (re.compile(r'(?<![a-zA-Z])ul(?![a-zA-Z])'),
             '微升应写μL（L大写）', 'μL', 'annotation'),
            (re.compile(r'ddH20|ddh20', re.IGNORECASE),
             '化学式ddH₂O中O是字母不是数字0', 'ddH₂O', 'annotation'),
        ]

        unit_spacing_pat = re.compile(
            r'(\d)(℃|°C|mol/L|mg/kg|μg|ng|pg|copies|CFU|IU|mmol|μmol)',
        )

        for i in range(start, min(end, len(paras))):
            text = paras[i].text
            if not text or len(text.strip()) < 3:
                continue
            section = self._get_section_label(i)
            loc_prefix = f'第{i+1}段 [{section}]'

            for pat, rule_desc, expected, source in unit_rules:
                for m in pat.finditer(text):
                    total_checks += 1
                    ctx_s = max(0, m.start() - 15)
                    ctx_e = min(len(text), m.end() + 15)
                    context = text[ctx_s:ctx_e]
                    self._add_issue(module, 'warning', loc_prefix, i,
                        paras[i].text.strip(),
                        rule_desc, expected, f'...{context}...', source)
                    error_count += 1

            for m in unit_spacing_pat.finditer(text):
                total_checks += 1
                ctx_s = max(0, m.start() - 10)
                ctx_e = min(len(text), m.end() + 10)
                self._add_issue(module, 'info', loc_prefix, i,
                    paras[i].text.strip(),
                    '数值与单位之间建议加空格（%除外）',
                    f'{m.group(1)} {m.group(2)}',
                    f'...{text[ctx_s:ctx_e]}...', 'supplement')

        total_checks = max(total_checks, 1)
        score = max(0, 7 * (1 - error_count / total_checks))
        self.scores[module] = (round(score, 1), 7)

    # --------------------------------------------------------
    # 检查模块 13: 内容规范
    # --------------------------------------------------------
    def check_content(self):
        """检查缩写全称、引用格式、拉丁学名、标点、摘要字数"""
        module = '内容规范'
        paras = self.doc.paragraphs
        error_count = 0
        total_checks = 0
        r = self.rules

        start = self.markers.get('introduction') or 0
        end = self.markers.get('references') or len(paras)

        # ---- 1. 中文摘要字数检查 ----
        total_checks += 1
        cn_idx = self.markers.get('abstract_cn')
        en_idx = self.markers.get('abstract_en')
        intro_idx = self.markers.get('introduction')
        if cn_idx is not None:
            abs_end = en_idx or intro_idx or (cn_idx + 80)
            cn_abstract_text = ''
            for j in range(cn_idx + 1, min(abs_end, len(paras))):
                text = paras[j].text.strip()
                if re.match(r'^(关\s*键\s*词|Keywords|Abstract)', text, re.IGNORECASE):
                    break
                if text and has_chinese(text):
                    cn_abstract_text += text

            cn_char_count = len(re.findall(r'[\u4e00-\u9fff]', cn_abstract_text))
            abs_min = r['abstract']['word_count_min']
            abs_max = r['abstract']['word_count_max']
            if cn_char_count < abs_min:
                self._add_issue(module, 'warning', '中文摘要', cn_idx,
                    f'摘要共{cn_char_count}字',
                    '硕士论文中文摘要约1000字', '约1000字',
                    f'{cn_char_count}字（偏少）', 'official')
                error_count += 0.5
            elif cn_char_count > abs_max:
                self._add_issue(module, 'warning', '中文摘要', cn_idx,
                    f'摘要共{cn_char_count}字',
                    '硕士论文中文摘要约1000字', '约1000字',
                    f'{cn_char_count}字（偏多）', 'official')
                error_count += 0.5

        # ---- 2. 缩写首次全称检查 ----
        total_checks += 1
        # 提取正文中所有大写缩写
        abbr_pat = re.compile(r'\b([A-Z]{2,}(?:-[A-Z0-9]+)?)\b')
        # 全称定义模式: "全称（Abbreviation，缩写）" 或 "Full Name (ABB)"
        define_pat = re.compile(r'[（(]\s*([A-Z]{2,})\s*[,，]?\s*([A-Z]{2,})?\s*[）)]')
        # 符号说明格式: "ABBR: Full name" or "ABBR：全称"
        symbol_def_pat = re.compile(r'^([A-Z][A-Z0-9\-]+)\s*[:：]')

        first_occurrences = {}  # abbr -> first para index
        defined_abbrs = set()

        # Extract abbreviations from 符号说明 section (before abstract)
        symbols_idx = self.markers.get('symbols')
        abstract_cn_idx = self.markers.get('abstract_cn')
        if symbols_idx is not None:
            symbols_end = abstract_cn_idx or (symbols_idx + 60)
            for j in range(symbols_idx + 1, min(symbols_end, len(paras))):
                sym_text = paras[j].text.strip()
                if not sym_text:
                    continue
                sym_m = symbol_def_pat.match(sym_text)
                if sym_m:
                    defined_abbrs.add(sym_m.group(1))
                # Also extract any uppercase abbreviations on the line
                for abb in abbr_pat.findall(sym_text):
                    defined_abbrs.add(abb)

        for i in range(start, min(end, len(paras))):
            text = paras[i].text
            if not text:
                continue

            # 记录定义的缩写
            for m in define_pat.finditer(text):
                defined_abbrs.add(m.group(1))
                if m.group(2):
                    defined_abbrs.add(m.group(2))

            # 也检查 "XXX, YYY" 模式在括号内
            bracket_content = re.findall(r'[（(]([^）)]+)[）)]', text)
            for bc in bracket_content:
                for abb in abbr_pat.findall(bc):
                    defined_abbrs.add(abb)

            # 记录首次出现
            for abb in abbr_pat.findall(text):
                if abb not in first_occurrences and len(abb) >= 3:
                    first_occurrences[abb] = i

        # 常见不需要解释的缩写
        common_abbrs = {'DNA', 'RNA', 'PCR', 'pH', 'UV', 'SDS', 'PAGE', 'EDTA',
                        'PBS', 'DMSO', 'Fig', 'USA', 'ALV', 'SPF', 'ELISA',
                        'qPCR', 'cDNA', 'mRNA', 'TAE', 'TBE', 'LB', 'OD',
                        'BLAST', 'NCBI', 'MEGA', 'RT', 'ORF', 'SYBR', 'MDA',
                        'MLV', 'HPRS', 'HRP', 'TMB', 'BSA', 'DEPC', 'DMEM',
                        'FBS', 'GAPDH', 'ANOVA', 'SD', 'CI', 'ROC', 'AUC',
                        'NTC', 'LOD', 'LOQ', 'CT', 'Ct', 'Tm'}

        undefined_abbrs = []
        for abb, first_idx in first_occurrences.items():
            if abb not in defined_abbrs and abb not in common_abbrs:
                if not any(abb in da for da in defined_abbrs):
                    undefined_abbrs.append((abb, first_idx))

        # 只报告最多8个
        for abb, idx in undefined_abbrs[:8]:
            self._add_issue(module, 'info', f'第{idx+1}段', idx,
                paras[idx].text.strip(),
                f'缩写"{abb}"首次出现时建议附全称',
                f'全称（{abb}）', f'直接使用{abb}', 'supplement')

        if len(undefined_abbrs) > 8:
            self._add_issue(module, 'info', '汇总', -1, '',
                f'共{len(undefined_abbrs)}个缩写未找到首次全称定义（仅展示前8个）',
                '', f'总计{len(undefined_abbrs)}个', 'supplement')

        # ---- 3. 参考文献正文引用格式 ----
        total_checks += 1
        # 检查正文中是否有 (作者, 年份) 或 (Author, Year) 格式引用
        cite_pat = re.compile(r'[（(]\s*[\u4e00-\u9fff\w]+.*?\d{4}\s*[）)]')
        cite_count = 0
        for i in range(start, min(end, len(paras))):
            text = paras[i].text
            if text:
                cite_count += len(cite_pat.findall(text))

        if cite_count < 10:
            self._add_issue(module, 'warning', '正文', -1, '',
                '正文中参考文献引用偏少（应在引用处标注"(作者，年份)"）',
                '多处引用标注', f'仅检测到约{cite_count}处', 'official')
            error_count += 0.5

        # ---- 4. 图表正文引用检查 ----
        total_checks += 1
        # 收集正文中引用的图表号
        body_text = ''
        for i in range(start, min(end, len(paras))):
            body_text += paras[i].text + '\n'

        fig_ref_pat = re.compile(r'图\s*(\d+)')
        tab_ref_pat = re.compile(r'表\s*(\d[\d\-\.]*)')
        referenced_figs = set(fig_ref_pat.findall(body_text))
        referenced_tabs = set(tab_ref_pat.findall(body_text))

        # 找正文中定义的图号
        defined_fig_nums = set()
        fig_cn_pat = re.compile(r'^图\s*(\d+)')
        for i in range(start, end):
            m = fig_cn_pat.match(paras[i].text.strip())
            if m:
                defined_fig_nums.add(m.group(1))

        unreferenced = defined_fig_nums - referenced_figs
        for fig_num in sorted(unreferenced, key=lambda x: int(x) if x.isdigit() else 0):
            self._add_issue(module, 'info', f'图{fig_num}', -1, '',
                f'图{fig_num}未在正文中被引用',
                '正文中应有"如图X所示"或"（图X）"', '未找到引用', 'supplement')

        # ---- 5. 拉丁学名斜体检查 ----
        total_checks += 1
        # 常见生物学名模式
        latin_pat = re.compile(
            r'\b(Escherichia\s+coli|Salmonella\s+\w+|Staphylococcus\s+\w+|'
            r'Streptococcus\s+\w+|Mycoplasma\s+\w+|Xanthomonas\s+\w+|'
            r'Pseudomonas\s+\w+|Gallus\s+gallus)\b'
        )
        non_italic_latin = 0
        for i in range(start, min(end, len(paras))):
            para = paras[i]
            for run in para.runs:
                if run.text and latin_pat.search(run.text):
                    if not run.font.italic:
                        non_italic_latin += 1
                        if non_italic_latin <= 5:
                            m = latin_pat.search(run.text)
                            self._add_issue(module, 'warning', f'第{i+1}段', i,
                                para.text.strip(),
                                f'拉丁学名"{m.group()}"应使用斜体',
                                '斜体', '正体', 'supplement')
                            error_count += 0.5

        if non_italic_latin > 5:
            self._add_issue(module, 'info', '汇总', -1, '',
                f'共{non_italic_latin}处拉丁学名未使用斜体（仅展示前5处）',
                '', f'总计{non_italic_latin}处', 'supplement')

        # ---- 6. 中文环境标点检查 ----
        total_checks += 1
        # 检查中文段落中出现英文标点
        half_punct_in_cn = 0
        en_punct_pat = re.compile(r'[\u4e00-\u9fff]\s*[,;:]\s*[\u4e00-\u9fff]')
        for i in range(start, min(end, len(paras))):
            text = paras[i].text
            if text and has_chinese(text):
                matches = en_punct_pat.findall(text)
                if matches:
                    half_punct_in_cn += len(matches)
                    if half_punct_in_cn <= 5:
                        m = en_punct_pat.search(text)
                        self._add_issue(module, 'warning', f'第{i+1}段', i,
                            text.strip(),
                            '中文语境中应使用全角标点（，；：）而非半角（,;:）',
                            '全角标点', f'检测到半角: ...{m.group()}...', 'supplement')
                        error_count += 0.5

        if half_punct_in_cn > 5:
            self._add_issue(module, 'info', '汇总', -1, '',
                f'共{half_punct_in_cn}处中文语境使用了半角标点（仅展示前5处）',
                '', f'总计{half_punct_in_cn}处', 'supplement')

        total_checks = max(total_checks, 1)
        score = max(0, 10 * (1 - error_count / total_checks))
        self.scores[module] = (round(score, 1), 10)

    # --------------------------------------------------------
    # 执行所有检查
    # --------------------------------------------------------
    def run_all_checks(self):
        """运行全部检查模块"""
        print('  [ 1/13] 检查页面设置...')
        self.check_page_setup()
        print('  [ 2/13] 检查封面...')
        self.check_cover()
        print('  [ 3/13] 检查摘要...')
        self.check_abstract()
        print('  [ 4/13] 检查目录...')
        self.check_toc()
        print('  [ 5/13] 检查正文格式...')
        self.check_body_text()
        print('  [ 6/13] 检查标题层级...')
        self.check_headings()
        print('  [ 7/13] 检查图表规范...')
        self.check_figures_tables()
        print('  [ 8/13] 检查页眉页脚...')
        self.check_headers_footers()
        print('  [ 9/13] 检查参考文献...')
        self.check_references()
        print('  [10/13] 检查结构完整性...')
        self.check_structure()
        print('  [11/13] 检查编号规范...')
        self.check_numbering()
        print('  [12/13] 检查单位符号...')
        self.check_units_symbols()
        print('  [13/13] 检查内容规范...')
        self.check_content()

    def get_total_score(self):
        return sum(s[0] for s in self.scores.values())

    def get_max_score(self):
        return sum(s[1] for s in self.scores.values())

    def get_report_data(self):
        """返回结构化报告数据（供 Web 应用调用），总分归一化为100分"""
        raw_total = self.get_total_score()
        raw_max = self.get_max_score()
        # 归一化到100分
        pct = (raw_total / raw_max * 100) if raw_max > 0 else 0
        score_100 = round(pct, 0)

        if pct >= 90: grade = 'A'
        elif pct >= 80: grade = 'B'
        elif pct >= 70: grade = 'C'
        elif pct >= 60: grade = 'D'
        else: grade = 'F'

        modules_order = ['页面设置', '封面', '摘要', '目录', '正文格式',
                         '标题层级', '图表规范', '页眉页脚', '参考文献', '其他规范',
                         '编号规范', '单位符号', '内容规范']
        modules = []
        for m in modules_order:
            earned, weight = self.scores.get(m, (0, 0))
            mod_pct = (earned / weight * 100) if weight > 0 else 0
            # 每个模块也归一化为该模块满分对应的100分比例
            norm_weight = round(weight / raw_max * 100, 1) if raw_max > 0 else 0
            norm_earned = round(earned / raw_max * 100, 1) if raw_max > 0 else 0
            issues_m = [i for i in self.issues if i.module == m]
            modules.append({
                'name': m,
                'earned': norm_earned,
                'weight': norm_weight,
                'pct': round(mod_pct, 1),
                'errors': sum(1 for i in issues_m if i.severity == 'error'),
                'warnings': sum(1 for i in issues_m if i.severity == 'warning'),
                'infos': sum(1 for i in issues_m if i.severity == 'info'),
            })

        issues = []
        for i in self.issues:
            issues.append({
                'module': i.module,
                'severity': i.severity,
                'severity_label': i.severity_label,
                'location': i.location,
                'para_index': i.para_index,
                'text_preview': i.text_preview,
                'rule': i.rule,
                'expected': i.expected,
                'actual': i.actual,
                'source': i.source,
                'source_label': i.source_label,
            })

        return {
            'filename': self.filename,
            'total_paras': self.total_paras,
            'total_tables': len(self.doc.tables),
            'total_images': len(self.doc.inline_shapes),
            'total_score': score_100,
            'max_score': 100,
            'pct': round(pct, 1),
            'grade': grade,
            'error_count': sum(1 for i in self.issues if i.severity == 'error'),
            'warning_count': sum(1 for i in self.issues if i.severity == 'warning'),
            'info_count': sum(1 for i in self.issues if i.severity == 'info'),
            'modules': modules,
            'issues': issues,
        }

    def _generate_module_filter_buttons(self):
        """生成模块筛选按钮 HTML"""
        modules_order = ['页面设置', '封面', '摘要', '目录', '正文格式',
                         '标题层级', '图表规范', '页眉页脚', '参考文献', '其他规范',
                         '编号规范', '单位符号', '内容规范']
        btns = []
        for m in modules_order:
            cnt = sum(1 for i in self.issues if i.module == m)
            if cnt > 0:
                btns.append(
                    f'<button class="filter-btn" onclick="filterBy(\'module\',\'{m}\')">'
                    f'{m} ({cnt})</button>')
        return '\n    '.join(btns)

    # --------------------------------------------------------
    # 生成 HTML 报告
    # --------------------------------------------------------
    def generate_html_report(self, output_path):
        raw_total = self.get_total_score()
        raw_max = self.get_max_score()
        # 归一化到100分
        pct = (raw_total / raw_max * 100) if raw_max > 0 else 0
        total = round(pct, 0)
        max_total = 100

        # 统计
        error_count = sum(1 for i in self.issues if i.severity == 'error')
        warning_count = sum(1 for i in self.issues if i.severity == 'warning')
        info_count = sum(1 for i in self.issues if i.severity == 'info')
        official_count = sum(1 for i in self.issues if i.source == 'official')
        supplement_count = sum(1 for i in self.issues if i.source == 'supplement')
        annotation_count = sum(1 for i in self.issues if i.source == 'annotation')

        # 评分等级和颜色
        if total >= 90:
            grade, grade_color = 'A (优秀)', '#22c55e'
        elif total >= 80:
            grade, grade_color = 'B (良好)', '#84cc16'
        elif total >= 70:
            grade, grade_color = 'C (中等)', '#eab308'
        elif total >= 60:
            grade, grade_color = 'D (及格)', '#f97316'
        else:
            grade, grade_color = 'F (不及格)', '#ef4444'

        # 各模块数据
        modules_order = ['页面设置', '封面', '摘要', '目录', '正文格式',
                         '标题层级', '图表规范', '页眉页脚', '参考文献', '其他规范',
                         '编号规范', '单位符号', '内容规范']

        module_rows = ''
        module_bars = ''
        for m in modules_order:
            earned, weight = self.scores.get(m, (0, 0))
            pct = (earned / weight * 100) if weight > 0 else 0
            issues_in_m = [i for i in self.issues if i.module == m]
            err = sum(1 for i in issues_in_m if i.severity == 'error')
            warn = sum(1 for i in issues_in_m if i.severity == 'warning')

            if pct >= 90:
                bar_color = '#22c55e'
            elif pct >= 70:
                bar_color = '#eab308'
            else:
                bar_color = '#ef4444'

            module_rows += f'''
            <tr>
                <td class="module-name">{m}</td>
                <td class="score-cell">{earned:.1f} / {weight}</td>
                <td>
                    <div class="bar-bg"><div class="bar-fill" style="width:{pct:.0f}%;background:{bar_color}"></div></div>
                </td>
                <td class="count-cell">{err}</td>
                <td class="count-cell">{warn}</td>
            </tr>'''

            module_bars += f'''
            <div class="radar-item">
                <div class="radar-label">{m}</div>
                <div class="radar-bar-bg">
                    <div class="radar-bar" style="height:{pct:.0f}%;background:{bar_color}"></div>
                </div>
                <div class="radar-score">{pct:.0f}%</div>
            </div>'''

        # 问题详情表
        issue_rows = ''
        for idx, issue in enumerate(self.issues):
            sev_class = {'error': 'sev-error', 'warning': 'sev-warning', 'info': 'sev-info'}[issue.severity]
            src_class = {'official': 'src-official', 'supplement': 'src-supplement',
                         'annotation': 'src-annotation'}.get(issue.source, 'src-supplement')

            issue_rows += f'''
            <tr class="issue-row {sev_class}-row" data-module="{issue.module}" data-severity="{issue.severity}" data-source="{issue.source}">
                <td><span class="badge {sev_class}">{issue.severity_label}</span></td>
                <td><span class="badge badge-module">{issue.module}</span></td>
                <td class="location-cell">{issue.location}</td>
                <td class="preview-cell" title="{issue.text_preview}">{issue.text_preview}</td>
                <td>{issue.rule}</td>
                <td class="expect-cell">{issue.expected}</td>
                <td class="actual-cell">{issue.actual}</td>
                <td><span class="badge {src_class}">{issue.source_label}</span></td>
            </tr>'''

        html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>硕士毕业论文格式审查报告</title>
<style>
:root {{
    --bg: #0f172a; --surface: #1e293b; --surface2: #334155;
    --text: #e2e8f0; --text2: #94a3b8; --text3: #64748b;
    --accent: #3b82f6; --red: #ef4444; --orange: #f97316;
    --yellow: #eab308; --green: #22c55e; --blue: #3b82f6;
}}
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', system-ui, sans-serif;
    background: var(--bg); color: var(--text); line-height: 1.6;
    padding: 24px; max-width: 1400px; margin: 0 auto;
}}
h1 {{ font-size: 1.8rem; font-weight: 700; margin-bottom: 4px; }}
h2 {{ font-size: 1.3rem; font-weight: 600; margin: 32px 0 16px; color: var(--text); }}
.subtitle {{ color: var(--text2); font-size: 0.9rem; margin-bottom: 24px; }}

/* 概览卡片 */
.overview {{ display: grid; grid-template-columns: 280px 1fr; gap: 24px; margin-bottom: 32px; }}
.score-card {{
    background: var(--surface); border-radius: 16px; padding: 32px;
    text-align: center; border: 1px solid var(--surface2);
}}
.score-big {{ font-size: 4rem; font-weight: 800; color: {grade_color}; line-height: 1; }}
.score-max {{ font-size: 1.2rem; color: var(--text3); }}
.grade {{ font-size: 1.4rem; font-weight: 600; color: {grade_color}; margin: 8px 0; }}
.stats {{ display: flex; justify-content: center; gap: 16px; margin-top: 16px; }}
.stat {{ text-align: center; }}
.stat-num {{ font-size: 1.5rem; font-weight: 700; }}
.stat-label {{ font-size: 0.75rem; color: var(--text3); }}

/* 柱状图 */
.chart-card {{
    background: var(--surface); border-radius: 16px; padding: 24px;
    border: 1px solid var(--surface2);
}}
.radar-container {{
    display: flex; align-items: flex-end; justify-content: space-around;
    height: 200px; padding: 0 8px;
}}
.radar-item {{ display: flex; flex-direction: column; align-items: center; flex: 1; }}
.radar-bar-bg {{
    width: 32px; height: 160px; background: var(--surface2); border-radius: 4px;
    position: relative; overflow: hidden;
}}
.radar-bar {{
    position: absolute; bottom: 0; width: 100%; border-radius: 4px 4px 0 0;
    transition: height 0.5s ease;
}}
.radar-label {{ font-size: 0.65rem; color: var(--text3); margin-top: 6px; text-align: center; writing-mode: horizontal-tb; }}
.radar-score {{ font-size: 0.75rem; font-weight: 600; color: var(--text2); margin-bottom: 4px; }}

/* 模块详情表 */
.module-table {{ width: 100%; border-collapse: collapse; margin-bottom: 24px; }}
.module-table th {{
    text-align: left; padding: 10px 12px; background: var(--surface2);
    font-size: 0.8rem; color: var(--text2); font-weight: 600;
}}
.module-table td {{ padding: 10px 12px; border-bottom: 1px solid var(--surface2); font-size: 0.9rem; }}
.module-name {{ font-weight: 600; }}
.score-cell {{ font-weight: 700; white-space: nowrap; }}
.count-cell {{ text-align: center; }}
.bar-bg {{ width: 100%; height: 8px; background: var(--surface2); border-radius: 4px; overflow: hidden; }}
.bar-fill {{ height: 100%; border-radius: 4px; transition: width 0.5s ease; }}

/* 筛选器 */
.filters {{ display: flex; gap: 8px; margin-bottom: 16px; flex-wrap: wrap; }}
.filter-btn {{
    padding: 6px 14px; border-radius: 20px; border: 1px solid var(--surface2);
    background: var(--surface); color: var(--text2); cursor: pointer;
    font-size: 0.8rem; transition: all 0.2s;
}}
.filter-btn:hover, .filter-btn.active {{
    background: var(--accent); color: white; border-color: var(--accent);
}}

/* 问题表 */
.issue-table {{ width: 100%; border-collapse: collapse; }}
.issue-table th {{
    text-align: left; padding: 10px 12px; background: var(--surface2);
    font-size: 0.75rem; color: var(--text2); font-weight: 600;
    position: sticky; top: 0; z-index: 10;
}}
.issue-table td {{ padding: 8px 12px; border-bottom: 1px solid rgba(51,65,85,0.5); font-size: 0.85rem; vertical-align: top; }}
.issue-row:hover {{ background: rgba(59,130,246,0.08); }}
.preview-cell {{ max-width: 180px; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; color: var(--text3); }}
.location-cell {{ white-space: nowrap; color: var(--text2); }}
.expect-cell {{ color: var(--green); }}
.actual-cell {{ color: var(--red); }}

/* 徽章 */
.badge {{
    display: inline-block; padding: 2px 10px; border-radius: 12px;
    font-size: 0.75rem; font-weight: 600; white-space: nowrap;
}}
.sev-error {{ background: rgba(239,68,68,0.15); color: #f87171; }}
.sev-warning {{ background: rgba(234,179,8,0.15); color: #facc15; }}
.sev-info {{ background: rgba(59,130,246,0.15); color: #60a5fa; }}
.src-official {{ background: rgba(139,92,246,0.15); color: #a78bfa; }}
.src-supplement {{ background: rgba(20,184,166,0.15); color: #2dd4bf; }}
.src-annotation {{ background: rgba(251,146,60,0.15); color: #fb923c; }}
.badge-module {{ background: var(--surface2); color: var(--text2); }}

.sev-error-row td:first-child {{ border-left: 3px solid var(--red); }}
.sev-warning-row td:first-child {{ border-left: 3px solid var(--yellow); }}
.sev-info-row td:first-child {{ border-left: 3px solid var(--blue); }}

/* 来源图例 */
.legend {{ display: flex; gap: 16px; margin-bottom: 16px; flex-wrap: wrap; }}
.legend-item {{ display: flex; align-items: center; gap: 6px; font-size: 0.8rem; color: var(--text2); }}

/* 打印 */
@media print {{
    body {{ background: white; color: #1e293b; padding: 12px; }}
    .score-card, .chart-card {{ border: 1px solid #e2e8f0; }}
    .filters {{ display: none; }}
    :root {{ --surface: #f8fafc; --surface2: #e2e8f0; --text: #1e293b; --text2: #64748b; --text3: #94a3b8; }}
    .score-big {{ color: #1e293b !important; }}
}}

/* 响应式 */
@media (max-width: 768px) {{
    .overview {{ grid-template-columns: 1fr; }}
    .radar-label {{ font-size: 0.55rem; }}
    .issue-table {{ font-size: 0.75rem; }}
}}
</style>
</head>
<body>

<h1>硕士毕业论文格式审查报告</h1>
<div class="subtitle">
    文件：{self.filename} &nbsp;|&nbsp; 段落数：{self.total_paras} &nbsp;|&nbsp;
    表格数：{len(self.doc.tables)} &nbsp;|&nbsp; 图片数：{len(self.doc.inline_shapes)}
</div>

<!-- 概览 -->
<div class="overview">
    <div class="score-card">
        <div class="score-big">{total:.0f}</div>
        <div class="score-max">/ {max_total}</div>
        <div class="grade">{grade}</div>
        <div class="stats">
            <div class="stat"><div class="stat-num" style="color:var(--red)">{error_count}</div><div class="stat-label">错误</div></div>
            <div class="stat"><div class="stat-num" style="color:var(--yellow)">{warning_count}</div><div class="stat-label">警告</div></div>
            <div class="stat"><div class="stat-num" style="color:var(--blue)">{info_count}</div><div class="stat-label">建议</div></div>
        </div>
        <div class="stats" style="margin-top:12px">
            <div class="stat"><div class="stat-num" style="color:#a78bfa">{official_count}</div><div class="stat-label">官方规定</div></div>
            <div class="stat"><div class="stat-num" style="color:#2dd4bf">{supplement_count}</div><div class="stat-label">专业补充</div></div>
            <div class="stat"><div class="stat-num" style="color:#fb923c">{annotation_count}</div><div class="stat-label">批注修订</div></div>
        </div>
    </div>
    <div class="chart-card">
        <h2 style="margin:0 0 12px">各模块得分率</h2>
        <div class="radar-container">{module_bars}</div>
    </div>
</div>

<!-- 模块评分 -->
<h2>模块评分详情</h2>
<table class="module-table">
    <thead><tr>
        <th>检查模块</th><th>得分</th><th>得分率</th><th>错误数</th><th>警告数</th>
    </tr></thead>
    <tbody>{module_rows}</tbody>
</table>

<!-- 来源图例 -->
<div class="legend">
    <div class="legend-item"><span class="badge src-official">官方规定</span> 学校文件明确要求的规则</div>
    <div class="legend-item"><span class="badge src-supplement">专业补充</span> 专业排版角度补充的通用细则</div>
    <div class="legend-item"><span class="badge src-annotation">批注修订</span> 从批注版修订意见中提炼的规则</div>
</div>

<!-- 筛选器 -->
<h2>问题详情（共 {len(self.issues)} 条）</h2>
<div class="filters" id="filters">
    <button class="filter-btn active" onclick="filterAll()">全部</button>
    <button class="filter-btn" onclick="filterBy('severity','error')" style="color:#f87171">错误 ({error_count})</button>
    <button class="filter-btn" onclick="filterBy('severity','warning')" style="color:#facc15">警告 ({warning_count})</button>
    <button class="filter-btn" onclick="filterBy('severity','info')" style="color:#60a5fa">建议 ({info_count})</button>
    <span style="color:var(--text3);padding:6px">|</span>
    <button class="filter-btn" onclick="filterBy('source','official')" style="color:#a78bfa">官方规定</button>
    <button class="filter-btn" onclick="filterBy('source','supplement')" style="color:#2dd4bf">专业补充</button>
    <button class="filter-btn" onclick="filterBy('source','annotation')" style="color:#fb923c">批注修订</button>
    <span style="color:var(--text3);padding:6px">|</span>
    {self._generate_module_filter_buttons()}
</div>

<!-- 问题表格 -->
<div style="overflow-x:auto; border-radius: 8px; border: 1px solid var(--surface2);">
<table class="issue-table">
    <thead><tr>
        <th>严重度</th><th>模块</th><th>位置</th><th>文本预览</th>
        <th>违反规则</th><th>期望值</th><th>实际值</th><th>规则来源</th>
    </tr></thead>
    <tbody>{issue_rows}</tbody>
</table>
</div>

<div style="text-align:center;color:var(--text3);font-size:0.75rem;margin-top:32px;padding:16px">
    论文格式一键体检 v1.0 &nbsp;|&nbsp; 审查标准基于学校官方规范 + 专业排版通用细则
</div>

<script>
function filterAll() {{
    document.querySelectorAll('.filter-btn').forEach(b => b.classList.remove('active'));
    event.target.classList.add('active');
    document.querySelectorAll('.issue-row').forEach(row => {{ row.style.display = ''; }});
}}
function filterBy(field, value) {{
    document.querySelectorAll('.filter-btn').forEach(b => b.classList.remove('active'));
    event.target.classList.add('active');
    document.querySelectorAll('.issue-row').forEach(row => {{
        row.style.display = row.dataset[field] === value ? '' : 'none';
    }});
}}
</script>

</body>
</html>'''

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)


# ============================================================
# 主入口
# ============================================================

def main():
    import argparse
    parser = argparse.ArgumentParser(description='硕士毕业论文格式审查工具')
    parser.add_argument('filepath', help='论文 .docx 文件路径')
    parser.add_argument('--title', '-t', help='指定论文中文题目（用于偶数页页眉校验）', default=None)
    args = parser.parse_args()

    filepath = args.filepath
    if not os.path.exists(filepath):
        print(f'错误: 文件不存在 - {filepath}')
        sys.exit(1)

    print(f'正在审查: {os.path.basename(filepath)}')
    if args.title:
        print(f'论文题目: {args.title}')
    print('=' * 50)

    checker = ThesisChecker(filepath, thesis_title=args.title)
    checker.run_all_checks()

    raw_total = checker.get_total_score()
    raw_max = checker.get_max_score()
    score_100 = round(raw_total / raw_max * 100) if raw_max > 0 else 0

    print('=' * 50)
    print(f'审查完成! 总分: {score_100} / 100')
    print(f'发现问题: {len(checker.issues)} 条')
    print(f'  - 错误: {sum(1 for i in checker.issues if i.severity == "error")} 条')
    print(f'  - 警告: {sum(1 for i in checker.issues if i.severity == "warning")} 条')
    print(f'  - 建议: {sum(1 for i in checker.issues if i.severity == "info")} 条')

    # 输出报告
    output_dir = os.path.dirname(os.path.abspath(filepath))
    output_path = os.path.join(output_dir, '格式审查报告.html')
    checker.generate_html_report(output_path)
    print(f'\n报告已生成: {output_path}')


if __name__ == '__main__':
    main()
